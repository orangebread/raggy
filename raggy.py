#!/usr/bin/env python3
"""Universal ChromaDB RAG Setup Script v2.0.0.

Enhanced with hybrid search, smart chunking, and normalized scoring.

Drop this into any project and run:
  python raggy.py build                       # Build/index all docs
  python raggy.py rebuild --fast              # Clean rebuild with faster model  
  python raggy.py search "your query"         # Semantic search with scores
  python raggy.py search "exact term" --hybrid # Hybrid semantic+keyword
  python raggy.py search "api" --expand        # Query expansion with synonyms
  python raggy.py interactive --quiet         # Interactive search mode
  python raggy.py status                      # Database stats with model info
  python raggy.py optimize                    # Benchmark and tune search modes
  python raggy.py search "query" --json       # Enhanced JSON output with scores

Key Features:
• Hybrid Search: Combines semantic + BM25 keyword ranking for exact matches
• Smart Chunking: Markdown-aware chunking preserving document structure  
• Normalized Scoring: 0-1 scores with human-readable labels
• Query Processing: Automatic expansion of domain terms
• Model Presets: --model-preset fast/balanced/multilingual/accurate
• Config Support: Optional raggy_config.yaml for customization
• Multilingual: Enhanced Dutch/English mixed content support
• Backward Compatible: All v1.x commands work unchanged
"""

# Standard library imports
import argparse
import glob
import hashlib
import importlib.util
import json
import math
import os
import re
import subprocess
import sys
import time
from datetime import datetime, timedelta
from collections import Counter, defaultdict
from pathlib import Path
from typing import (
    Any,
    Callable,
    Dict,
    Iterator,
    List,
    Optional,
    Set,
    Tuple,
    Union,
)

# Version information
__version__ = "2.0.0"

# Constants
CHUNK_READ_SIZE = 8192  # 8KB chunks for file reading
MAX_CACHE_SIZE = 1000   # Maximum number of cached embeddings  
CACHE_TTL = 3600       # Cache time-to-live in seconds (1 hour)
MAX_FILE_SIZE_MB = 100  # Maximum file size in MB
SESSION_CACHE_HOURS = 24  # Hours before update check
UPDATE_TIMEOUT_SECONDS = 2  # API timeout for update checks
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 200
DEFAULT_RESULTS = 5
DEFAULT_CONTEXT_CHARS = 200
DEFAULT_HYBRID_WEIGHT = 0.7
MANIFEST_FILENAME = "index-manifest.json"
MANIFEST_SCHEMA_VERSION = 1
DEFAULT_SOFT_CHUNK_LIMIT = 8000
DEFAULT_HARD_CHUNK_LIMIT = 12000
DEFAULT_SOFT_DOCUMENT_LIMIT = 300
DEFAULT_PER_DOC_CHUNK_LIMIT = 800
DEFAULT_HOT_RETENTION_DAYS = 21
DEFAULT_MIN_HOT_UPDATES = 5
DEFAULT_ARCHIVE_SUBDIR = "archive/development_state"
DEFAULT_DIGEST_SUBDIR = "summaries"

# File type constants
SUPPORTED_EXTENSIONS = [".md", ".pdf", ".docx", ".txt"]
GLOB_PATTERNS = ["**/*.md", "**/*.pdf", "**/*.docx", "**/*.txt"]

# Model presets
FAST_MODEL = "paraphrase-MiniLM-L3-v2"
DEFAULT_MODEL = "all-MiniLM-L6-v2"
MULTILINGUAL_MODEL = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
ACCURATE_MODEL = "all-mpnet-base-v2"

# Configure UTF-8 encoding for cross-platform compatibility
if sys.platform == "win32":
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    # Try to configure console for UTF-8 on Windows
    try:
        import codecs

        if hasattr(sys.stdout, "reconfigure"):
            sys.stdout.reconfigure(encoding="utf-8")
        if hasattr(sys.stderr, "reconfigure"):
            sys.stderr.reconfigure(encoding="utf-8")
    except (AttributeError, OSError):
        pass  # Ignore encoding configuration errors


# Cross-platform emoji/symbol support
def get_symbols() -> Dict[str, str]:
    """Get appropriate symbols based on platform/terminal support."""
    try:
        # Test if terminal supports unicode
        test = "🔍"
        print(test, end="")
        print("\b \b", end="")  # backspace and clear
        return {
            "search": "🔍", 
            "found": "📋", 
            "success": "✅", 
            "bye": "👋"
        }
    except UnicodeEncodeError:
        return {
            "search": "[Search]",
            "found": "[Found]", 
            "success": "[Success]",
            "bye": "[Bye]",
        }


SYMBOLS = get_symbols()

# Pre-compiled regex patterns for performance
WORD_PATTERN = re.compile(r"\b\w+\b")
NEGATIVE_TERM_PATTERN = re.compile(r"-\w+")
AND_TERM_PATTERN = re.compile(r"\w+(?=\s+AND)", re.IGNORECASE)
QUOTED_PHRASE_PATTERN = re.compile(r'"([^"]+)"')
HEADER_PATTERN = re.compile(r"(^#{1,6}\s+.*$)", re.MULTILINE)
SENTENCE_BOUNDARY_PATTERN = re.compile(r"[.!?\n]")
WINDOWS_PATH_PATTERN = re.compile(r'[A-Za-z]:[\\\/][^\\\/\s]*[\\\/]')
UNIX_PATH_PATTERN = re.compile(r'\/[^\/\s]*\/')
FILE_URL_PATTERN = re.compile(r'\bfile:\/\/[^\s]*')


def validate_path(file_path: Path, base_path: Optional[Path] = None) -> bool:
    """Validate file path to prevent directory traversal attacks."""
    try:
        # Resolve the path to get absolute path
        resolved_path = file_path.resolve()
        
        if base_path is None:
            base_path = Path.cwd()
        else:
            base_path = base_path.resolve()
        
        # Check if the resolved path is within the base directory
        try:
            resolved_path.relative_to(base_path)
            return True
        except ValueError:
            # Path is outside the base directory
            return False
    except (OSError, ValueError):
        return False


def sanitize_error_message(error_msg: str) -> str:
    """Sanitize error messages to prevent information leakage."""
    # Remove potentially sensitive path information using pre-compiled patterns
    sanitized = WINDOWS_PATH_PATTERN.sub('', error_msg)  # Windows paths
    sanitized = UNIX_PATH_PATTERN.sub('/', sanitized)  # Unix paths  
    sanitized = FILE_URL_PATTERN.sub('[FILE_PATH]', sanitized)
    return sanitized


def log_error(message: str, error: Optional[Exception] = None, *, quiet: bool = False) -> None:
    """Centralized error logging with consistent formatting."""
    if quiet:
        return
    
    if error:
        sanitized_error = sanitize_error_message(str(error))
        print(f"ERROR: {message}: {sanitized_error}")
    else:
        print(f"ERROR: {message}")


def log_warning(message: str, error: Optional[Exception] = None, *, quiet: bool = False) -> None:
    """Centralized warning logging with consistent formatting."""
    if quiet:
        return
    
    if error:
        sanitized_error = sanitize_error_message(str(error))
        print(f"Warning: {message}: {sanitized_error}")
    else:
        print(f"Warning: {message}")


def handle_file_error(file_path: Path, operation: str, error: Exception, *, quiet: bool = False) -> None:
    """Standardized file operation error handling."""
    if isinstance(error, (FileNotFoundError, PermissionError)):
        log_error(f"Cannot {operation} {file_path.name} - {type(error).__name__}", quiet=quiet)
    elif isinstance(error, UnicodeDecodeError):
        log_error(f"Cannot {operation} {file_path.name} - encoding issue", quiet=quiet)
    else:
        log_error(f"Cannot {operation} {file_path.name}", error, quiet=quiet)


def _default_manifest() -> Dict[str, Any]:
    """Return an empty manifest structure."""
    return {
        "schema_version": MANIFEST_SCHEMA_VERSION,
        "model_name": None,
        "embedding_dim": None,
        "documents": {},
        "last_built_at": None,
    }


def load_manifest(manifest_path: Path, *, quiet: bool = False) -> Dict[str, Any]:
    """Load the index manifest if it exists, otherwise return defaults."""
    if not manifest_path.exists():
        return _default_manifest()

    try:
        with open(manifest_path, "r", encoding="utf-8") as manifest_file:
            data = json.load(manifest_file)
    except (json.JSONDecodeError, OSError) as err:
        log_warning(
            f"Could not read manifest at {manifest_path.name}; starting fresh",
            err,
            quiet=quiet,
        )
        return _default_manifest()

    data.setdefault("schema_version", MANIFEST_SCHEMA_VERSION)
    data.setdefault("model_name", None)
    data.setdefault("embedding_dim", None)
    data.setdefault("documents", {})
    data.setdefault("last_built_at", None)
    return data


def save_manifest(manifest_path: Path, manifest: Dict[str, Any]) -> None:
    """Persist the index manifest to disk."""
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest["schema_version"] = MANIFEST_SCHEMA_VERSION
    with open(manifest_path, "w", encoding="utf-8") as manifest_file:
        json.dump(manifest, manifest_file, indent=2, sort_keys=True)


def check_for_updates(
    quiet: bool = False, config: Optional[Dict[str, Any]] = None
) -> None:
    """Check GitHub for latest version once per session (non-intrusive)."""
    if quiet:
        return
    
    # Load configuration for update settings
    if config is None:
        config = {}
    
    updates_config = config.get("updates", {})
    if not updates_config.get("check_enabled", True):
        return
    
    # Use configured repo or default placeholder
    github_repo = updates_config.get("github_repo", "dimitritholen/raggy")
    
    # Session tracking to avoid frequent checks
    session_file = Path.home() / ".raggy_session"
    
    # Check if already checked in last 24 hours
    if session_file.exists():
        try:
            cache_age = time.time() - session_file.stat().st_mtime
            if cache_age < SESSION_CACHE_HOURS * 3600:  # 24 hours
                return
        except (OSError, AttributeError):
            pass  # If we can't check file time, proceed with check
    
    try:
        # Import urllib only when needed to avoid startup cost
        import urllib.request
        import urllib.error
        
        # Quick timeout to not delay startup
        api_url = f"https://api.github.com/repos/{github_repo}/releases/latest"
        
        with urllib.request.urlopen(api_url, timeout=UPDATE_TIMEOUT_SECONDS) as response:
            if response.status == 200:
                data = json.loads(response.read().decode('utf-8'))
                latest_version = data.get("tag_name", "").lstrip("v")
                
                if latest_version and latest_version != __version__:
                    # Use HTML URL from response or construct fallback
                    github_url = data.get("html_url")
                    if not github_url:
                        base_url = f"https://github.com/{github_repo}"
                        github_url = f"{base_url}/releases/latest"
                    
                    print(f"📦 Raggy update available: v{latest_version} → {github_url}")
        
        # Update session file to mark check as done
        try:
            session_file.touch()
        except (OSError, PermissionError):
            pass  # If we can't create session file, just skip tracking
            
    except (
        urllib.error.URLError, 
        urllib.error.HTTPError, 
        json.JSONDecodeError,
        ConnectionError, 
        TimeoutError, 
        Exception
    ):
        # Silently fail - don't interrupt user workflow with network issues
        # This includes any import errors, network timeouts, or API issues
        pass


class BM25Scorer:
    """Lightweight BM25 implementation for keyword scoring."""

    def __init__(self, k1: float = 1.2, b: float = 0.75) -> None:
        self.k1 = k1
        self.b = b
        self.doc_lengths: List[int] = []
        self.avg_doc_length = 0.0
        self.doc_count = 0
        self.term_frequencies: List[Dict[str, int]] = []
        self.idf_scores: Dict[str, float] = {}

    def fit(self, documents: List[str]) -> None:
        """Build BM25 index from documents."""
        self.doc_count = len(documents)
        self.doc_lengths = []
        self.term_frequencies = []
        doc_term_counts: Dict[str, int] = defaultdict(int)

        # Calculate term frequencies and document lengths
        for doc in documents:
            terms = self._tokenize(doc)
            self.doc_lengths.append(len(terms))

            term_freq = Counter(terms)
            self.term_frequencies.append(term_freq)

            # Count documents containing each term
            for term in set(terms):
                doc_term_counts[term] += 1

        self.avg_doc_length = (
            sum(self.doc_lengths) / len(self.doc_lengths) 
            if self.doc_lengths else 0.0
        )

        # Calculate IDF scores
        for term, doc_freq in doc_term_counts.items():
            # Use standard BM25 IDF: log((N + 1) / df)
            # This avoids negative scores and is more stable for small datasets
            self.idf_scores[term] = math.log((self.doc_count + 1) / doc_freq)

    def score(self, query: str, doc_index: int) -> float:
        """Calculate BM25 score for query against document."""
        if doc_index >= len(self.term_frequencies):
            return 0.0

        query_terms = self._tokenize(query)
        score = 0.0
        doc_length = self.doc_lengths[doc_index]
        term_freq = self.term_frequencies[doc_index]

        for term in query_terms:
            if term in term_freq:
                tf = term_freq[term]
                idf = self.idf_scores.get(term, 0.0)

                numerator = tf * (self.k1 + 1)
                length_normalization = (
                    1 - self.b + self.b * (doc_length / self.avg_doc_length)
                )
                denominator = tf + self.k1 * length_normalization
                score += idf * (numerator / denominator)

        return max(0.0, score)  # Ensure non-negative scores

    def _tokenize(self, text: str) -> List[str]:
        """Simple tokenization"""
        # Convert to lowercase and extract alphanumeric sequences using pre-compiled pattern
        return WORD_PATTERN.findall(text.lower())


class QueryProcessor:
    """Enhanced query processing with expansion and operators."""

    def __init__(
        self, custom_expansions: Optional[Dict[str, List[str]]] = None
    ) -> None:
        # Default expansions - can be overridden via config
        self.expansions = custom_expansions or {
            # Common technical terms
            "api": ["api", "application programming interface"],
            "ml": ["ml", "machine learning"],
            "ai": ["ai", "artificial intelligence"],
            "ui": ["ui", "user interface"],
            "ux": ["ux", "user experience"],
            # Can be extended via configuration file
        }

    def process(self, query: str) -> Dict[str, Any]:
        """Process query and return enhanced version with metadata."""
        original = query.strip()

        # Detect query type
        query_type = self._detect_type(original)

        # Handle exact phrase queries (quoted)
        if query_type == "exact":
            phrase = QUOTED_PHRASE_PATTERN.findall(original)[0]
            return {
                "processed": phrase,
                "original": original,
                "type": "exact",
                "boost_exact": True,
                "terms": [phrase],
            }

        # Expand terms
        expanded = self._expand_query(original)

        # Extract boolean operators
        must_have, must_not = self._extract_operators(expanded)

        return {
            "processed": expanded,
            "original": original,
            "type": query_type,
            "boost_exact": False,
            "must_have": must_have,
            "must_not": must_not,
            "terms": WORD_PATTERN.findall(expanded.lower()),
        }

    def _detect_type(self, query: str) -> str:
        """Detect query type."""
        if '"' in query:
            return "exact"
        
        question_words = ["how", "what", "why", "when", "where", "who"]
        if any(word in query.lower() for word in question_words):
            return "question"
        
        boolean_operators = [" AND ", " OR ", " -"]
        query_upper = query.upper()
        if any(op in query_upper or op.strip() in query for op in boolean_operators):
            return "boolean"
        
        return "keyword"

    def _expand_query(self, query: str) -> str:
        """Expand query with synonyms"""
        expanded = query.lower()
        for term, expansions in self.expansions.items():
            if term in expanded:
                # Add expansions as OR terms
                expansion_str = " OR ".join(expansions[1:])  # Skip the original term
                if expansion_str:
                    expanded = expanded.replace(term, f"({term} OR {expansion_str})")
        return expanded

    def _extract_operators(self, query: str) -> Tuple[List[str], List[str]]:
        """Extract boolean operators"""
        must_have = []
        must_not = []

        # Extract negative terms (preceded by -)
        negative_terms = NEGATIVE_TERM_PATTERN.findall(query)
        for term in negative_terms:
            must_not.append(term[1:])  # Remove the -

        # Extract AND terms
        and_terms = AND_TERM_PATTERN.findall(query)
        must_have.extend(and_terms)

        return must_have, must_not


# Scoring utility functions (converted from class with static methods)
def normalize_cosine_distance(distance: float) -> float:
    """Convert cosine distance to similarity (0-1, higher is better)."""
    # ChromaDB returns cosine distance (0 = identical, 2 = opposite)
    # Convert to similarity: 1 - (distance / 2)
    similarity = max(0.0, min(1.0, 1.0 - (distance / 2.0)))
    return similarity


def normalize_hybrid_score(
    semantic_score: float, 
    keyword_score: float, 
    semantic_weight: float = DEFAULT_HYBRID_WEIGHT
) -> float:
    """Combine semantic and keyword scores."""
    # Normalize keyword score to 0-1 range (assuming max BM25 ~10)
    normalized_keyword = min(1.0, keyword_score / 10.0)

    weighted_semantic = semantic_weight * semantic_score
    weighted_keyword = (1 - semantic_weight) * normalized_keyword
    return weighted_semantic + weighted_keyword


def interpret_score(score: float) -> str:
    """Provide human-readable score interpretation."""
    if score >= 0.8:
        return "Excellent"
    elif score >= 0.6:
        return "Good"
    elif score >= 0.4:
        return "Fair"
    else:
        return "Poor"


def load_config(config_path: Optional[str] = None) -> Dict[str, Any]:
    """Load optional configuration file."""
    default_config = {
        "search": {
            "hybrid_weight": DEFAULT_HYBRID_WEIGHT,
            "chunk_size": DEFAULT_CHUNK_SIZE,
            "chunk_overlap": DEFAULT_CHUNK_OVERLAP,
            "rerank": True,
            "show_scores": True,
            "context_chars": DEFAULT_CONTEXT_CHARS,
            "max_results": DEFAULT_RESULTS,
            "expansions": {
                # Add domain-specific expansions here
                "api": ["api", "application programming interface"],
                "ml": ["ml", "machine learning"],
                "ai": ["ai", "artificial intelligence"],
                "ui": ["ui", "user interface"],
                "ux": ["ux", "user experience"],
            },
        },
        "models": {
            "default": DEFAULT_MODEL,
            "fast": FAST_MODEL,
            "multilingual": MULTILINGUAL_MODEL,
            "accurate": ACCURATE_MODEL,
        },
        "chunking": {
            "smart": True,  # Enable by default for better documentation structure
            "preserve_headers": True,
            "min_chunk_size": 300,
            "max_chunk_size": 1500,
        },
        "maintenance": {
            "thresholds": {
                "soft_chunk_limit": DEFAULT_SOFT_CHUNK_LIMIT,
                "hard_chunk_limit": DEFAULT_HARD_CHUNK_LIMIT,
                "soft_document_limit": DEFAULT_SOFT_DOCUMENT_LIMIT,
                "per_document_chunk_limit": DEFAULT_PER_DOC_CHUNK_LIMIT,
            },
            "retention": {
                "hot_days": DEFAULT_HOT_RETENTION_DAYS,
                "min_hot_updates": DEFAULT_MIN_HOT_UPDATES,
            },
            "paths": {
                "archive_dir": DEFAULT_ARCHIVE_SUBDIR,
                "digest_dir": DEFAULT_DIGEST_SUBDIR,
                "compaction_source": "CHANGELOG.md",  # File to compact (legacy: DEVELOPMENT_STATE.md)
            },
            "auto_compact": {
                "rebuild_after_compact": True,
            },
        },
        "updates": {
            "check_enabled": True,  # Enable update checking by default
            "github_repo": "dimitritholen/raggy",  # Repository for update checks
        },
    }

    # Try to load config file
    config_file = Path(config_path or "raggy_config.yaml")
    if config_file.exists():
        try:
            import yaml

            with open(config_file, "r") as f:
                user_config = yaml.safe_load(f)

            # Merge with defaults
            _merge_configs(default_config, user_config)
        except ImportError:
            log_warning("PyYAML not installed, using default config", quiet=False)
        except (FileNotFoundError, PermissionError) as e:
            log_warning(f"Could not access config file {config_file}", e, quiet=False)
        except Exception as yaml_error:
            # Handle YAML parsing errors (yaml module imported locally)
            if any(keyword in str(yaml_error).lower() 
                   for keyword in ["yaml", "scanner", "parser", "constructor"]):
                log_warning(f"Invalid YAML format in {config_file}", yaml_error, quiet=False)
            else:
                # Re-raise if it's not a YAML parsing error
                raise yaml_error
        except Exception as e:
            log_warning(f"Unexpected error loading config file {config_file}", e, quiet=False)

    return default_config


def _merge_configs(default: Dict[str, Any], user: Dict[str, Any]) -> None:
    """Recursively merge user config into default config."""
    for key, value in user.items():
        if (
            key in default
            and isinstance(default[key], dict)
            and isinstance(value, dict)
        ):
            _merge_configs(default[key], value)
        else:
            default[key] = value


def get_cache_file():
    """Get path for dependency cache file"""
    return Path.cwd() / ".raggy_deps_cache.json"


def load_deps_cache():
    """Load dependency cache"""
    cache_file = get_cache_file()
    if cache_file.exists():
        try:
            with open(cache_file, "r") as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError, PermissionError):
            pass
    return {}


def save_deps_cache(cache):
    """Save dependency cache"""
    cache_file = get_cache_file()
    try:
        with open(cache_file, "w") as f:
            json.dump(cache, f, indent=2)
    except (PermissionError, OSError):
        pass  # Ignore cache save errors


def check_uv_available():
    """Check if uv is available"""
    try:
        subprocess.check_call(
            ["uv", "--version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("ERROR: uv is not available or not in PATH")
        print(
            "Please install uv first: https://docs.astral.sh/uv/getting-started/installation/"
        )
        print("Or run: curl -LsSf https://astral.sh/uv/install.sh | sh")
        return False


def check_environment_setup():
    """Check if project environment is properly set up"""
    venv_path = Path(".venv")
    pyproject_path = Path("pyproject.toml")
    
    if not venv_path.exists():
        return False, "virtual_environment"
    
    if not pyproject_path.exists():
        return False, "pyproject"
    
    # Check if virtual environment is activated or can be used
    try:
        # Check if we can run python in the venv
        if sys.platform == "win32":
            python_exe = venv_path / "Scripts" / "python.exe"
        else:
            python_exe = venv_path / "bin" / "python"
        
        if not python_exe.exists():
            return False, "invalid_venv"
    except (OSError, AttributeError):
        return False, "invalid_venv"
    
    return True, "ok"


def _create_virtual_environment(quiet: bool = False) -> bool:
    """Create virtual environment if it doesn't exist."""
    venv_path = Path(".venv")
    if not venv_path.exists():
        if not quiet:
            print("Creating virtual environment...")
        try:
            subprocess.check_call(
                ["uv", "venv"], 
                stdout=subprocess.DEVNULL if quiet else None
            )
        except subprocess.CalledProcessError as e:
            print(f"ERROR: Failed to create virtual environment: {e}")
            return False
    return True


def _create_project_config(quiet: bool = False) -> bool:
    """Create minimal pyproject.toml if it doesn't exist."""
    pyproject_path = Path("pyproject.toml")
    if not pyproject_path.exists():
        if not quiet:
            print("Creating pyproject.toml...")
        
        pyproject_content = """[project]
name = "raggy-project"
version = "0.1.0"
description = "RAG project using Universal ChromaDB RAG System"
requires-python = ">=3.8"
dependencies = [
    "chromadb>=0.4.0",
    "sentence-transformers>=2.2.0",
    "PyPDF2>=3.0.0",
    "python-docx>=1.0.0",
]

[project.optional-dependencies]
magic-win = ["python-magic-bin>=0.4.14"]
magic-unix = ["python-magic"]

[build-system]
requires = ["hatchling"]
build-backend = "hatchling.build"
"""
        
        try:
            with open(pyproject_path, "w", encoding="utf-8") as f:
                f.write(pyproject_content)
        except (PermissionError, OSError) as e:
            log_error("Failed to create pyproject.toml", e, quiet=False)
            return False
    return True


def _install_dependencies(quiet: bool = False) -> bool:
    """Install core and platform-specific dependencies."""
    if not quiet:
        print("Installing dependencies...")
    
    try:
        # Install base dependencies
        base_deps = [
            "chromadb>=0.4.0", 
            "sentence-transformers>=2.2.0", 
            "PyPDF2>=3.0.0",
            "python-docx>=1.0.0"
        ]
        subprocess.check_call(
            ["uv", "pip", "install"] + base_deps,
            stdout=subprocess.DEVNULL if quiet else None
        )
        
        # Install platform-specific magic library
        _install_magic_library(quiet)
        
    except subprocess.CalledProcessError as e:
        print(f"ERROR: Failed to install dependencies: {e}")
        print("Manual install: uv pip install chromadb sentence-transformers PyPDF2")
        return False
    
    return True


def _install_magic_library(quiet: bool = False) -> None:
    """Install platform-specific magic library for file type detection."""
    magic_package = (
        "python-magic-bin>=0.4.14" if sys.platform == "win32" 
        else "python-magic"
    )
    
    try:
        subprocess.check_call(
            ["uv", "pip", "install", magic_package],
            stdout=subprocess.DEVNULL if quiet else None
        )
    except subprocess.CalledProcessError:
        if not quiet:
            package_name = magic_package.split(">")[0]  # Remove version spec
            warning = f"Warning: Could not install {package_name}. "
            warning += "File type detection may be limited."
            print(warning)


def _create_docs_directory(quiet: bool = False) -> Optional[Path]:
    """Create docs directory if it doesn't exist."""
    docs_path = Path("docs")
    if not docs_path.exists():
        try:
            docs_path.mkdir()
            if not quiet:
                print("Created docs/ directory - add your documentation files here")
        except OSError as e:
            print(f"ERROR: Failed to create docs directory: {e}")
            return None
    return docs_path


def _create_development_state_file(docs_path: Path, quiet: bool = False) -> bool:
    """Create initial DEVELOPMENT_STATE.md for AI workflow tracking."""
    dev_state_path = docs_path / "DEVELOPMENT_STATE.md"
    if not dev_state_path.exists():
        if not quiet:
            print("Creating initial DEVELOPMENT_STATE.md...")
        
        timestamp = time.strftime('%Y-%m-%d %H:%M:%S')
        dev_state_content = f"""# Development State

**Last Updated:** {timestamp}
**RAG System:** Raggy v2.0.0 - Universal ChromaDB RAG Setup

## Project Status: INITIALIZED

### COMPLETED:
- ✅ Raggy environment initialized with `python raggy.py init`
- ✅ Virtual environment (.venv) created and activated
- ✅ Dependencies installed: chromadb, sentence-transformers, PyPDF2, python-docx
- ✅ Project configuration (pyproject.toml) generated
- ✅ Example configuration (raggy_config_example.yaml) created
- ✅ Documentation directory (docs/) established
- ✅ Development state tracking initialized

### CURRENT SETUP:
- **Supported formats:** .md (Markdown), .pdf (PDF), .docx (Word), .txt (Plain text)
- **Search modes:** Semantic, Hybrid (semantic + keyword), Query expansion
- **Model presets:** fast/balanced/multilingual/accurate
- **Local processing:** 100% offline, zero API costs

### NEXT STEPS:
1. **Add documentation files** to the docs/ directory
2. **Optional:** Copy raggy_config_example.yaml to raggy_config.yaml and customize expansions
3. **Build the RAG database:** Run `python raggy.py build`
4. **Test search functionality:** Run `python raggy.py search "your query"`
5. **Configure AI agents** with the knowledge-driven workflow from README.md

### DECISIONS:
- Chose raggy for zero-cost local RAG implementation
- Configured for multi-format document support (.md, .pdf, .docx, .txt)
- Set up for AI agent integration with continuous development state tracking

### ARCHITECTURE:
- **Vector Database:** ChromaDB (local storage in ./vectordb/)
- **Embeddings:** sentence-transformers (local, no API costs)
- **Search Engine:** Hybrid semantic + BM25 keyword ranking
- **Document Processing:** Smart chunking with markdown awareness

### BLOCKERS:
- None - system ready for document ingestion and usage

---

*This file tracks development progress for AI agent continuity. Update after each significant task or decision.*
"""
        
        try:
            with open(dev_state_path, "w", encoding="utf-8") as f:
                f.write(dev_state_content)
        except (PermissionError, OSError) as e:
            log_warning("Could not create DEVELOPMENT_STATE.md", e, quiet=False)
            return False
    
    return True


def setup_environment(quiet: bool = False) -> bool:
    """Set up the project environment from scratch."""
    if not quiet:
        print("🚀 Setting up raggy environment...")
    
    # Check if uv is available
    if not check_uv_available():
        return False
    
    # Create virtual environment
    if not _create_virtual_environment(quiet):
        return False
    
    # Create minimal pyproject.toml if it doesn't exist
    if not _create_project_config(quiet):
        return False
    
    # Install dependencies
    if not _install_dependencies(quiet):
        return False
    
    # Create docs directory if it doesn't exist
    docs_path = _create_docs_directory(quiet)
    if docs_path is None:
        return False
    
    # Create initial DEVELOPMENT_STATE.md for AI workflow tracking
    if not _create_development_state_file(docs_path, quiet):
        # Warning already printed, continue anyway
        pass
    
    # Create example config file if it doesn't exist
    if not _create_example_config(quiet):
        # Warning already printed, continue anyway
        pass
    
    if not quiet:
        _print_setup_summary()
    
    return True


def _create_example_config(quiet: bool = False) -> bool:
    """Create example configuration file, but only if no real config exists."""
    config_example_path = Path("raggy_config_example.yaml")
    real_config_path = Path("raggy_config.yaml")

    # If a real config exists, do not create the example file
    if real_config_path.exists():
        if not quiet:
            print("raggy_config.yaml exists; skipping example config generation")
        return True

    if not config_example_path.exists():
        if not quiet:
            print("Creating raggy_config_example.yaml...")
        
        config_content = """# raggy_config_example.yaml - Example Configuration File
# Copy this to raggy_config.yaml and customize for your domain

search:
  hybrid_weight: 0.7  # Balance between semantic (0.7) and keyword (0.3) search
  chunk_size: 1000
  chunk_overlap: 200
  rerank: true
  show_scores: true
  context_chars: 200
  max_results: 5
  
  # Domain-specific query expansions
  # Add your own terms here for automatic expansion
  expansions:
    # Technical terms (examples)
    api: ["api", "application programming interface", "rest api", "web service"]
    ml: ["ml", "machine learning", "artificial intelligence"]
    ai: ["ai", "artificial intelligence", "machine learning"]
    ui: ["ui", "user interface", "frontend", "user experience"]
    ux: ["ux", "user experience", "usability", "user interface"]
    
    # Business terms (examples)
    roi: ["roi", "return on investment", "profitability"]
    kpi: ["kpi", "key performance indicator", "metrics"]
    crm: ["crm", "customer relationship management", "customer management"]
    
    # Development terms (examples)
    ci: ["ci", "continuous integration", "build automation"]
    cd: ["cd", "continuous deployment", "continuous delivery"]
    devops: ["devops", "development operations", "infrastructure"]
    
    # Add your domain-specific terms here:
    # mycompany: ["mycompany", "company name", "organization"]
    # myproduct: ["myproduct", "product name", "solution"]

models:
  default: "all-MiniLM-L6-v2"           # Balanced speed/accuracy
  fast: "paraphrase-MiniLM-L3-v2"       # Fastest, smaller model  
  multilingual: "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"  # Multi-language
  accurate: "all-mpnet-base-v2"         # Best accuracy, slower

chunking:
  smart: true           # Enable markdown-aware smart chunking
  preserve_headers: true # Include section headers in chunks
  min_chunk_size: 300   # Minimum chunk size in characters
  max_chunk_size: 1500  # Maximum chunk size in characters

maintenance:
  thresholds:
    soft_chunk_limit: 8000      # Warn when active chunks exceed this value
    hard_chunk_limit: 12000     # Require compaction once this limit is reached
    soft_document_limit: 300    # Warn when document count exceeds this value
    per_document_chunk_limit: 800 # Flag any single doc that exceeds this many chunks
  retention:
    hot_days: 21                # Keep updates newer than this many days in the hot ledger
    min_hot_updates: 5          # Always retain at least this many recent updates
  paths:
    archive_dir: "archive/development_state" # Relative to docs/ directory
    digest_dir: "summaries"                  # Monthly digest output directory
  auto_compact:
    rebuild_after_compact: true  # Automatically rebuild index after compaction

# Usage:
# 1. Copy this file to raggy_config.yaml  
# 2. Customize the expansions section with your domain terms
# 3. Adjust model and chunking settings as needed
# 4. Run: python raggy.py search "your-term" --expand
"""
        
        try:
            with open(config_example_path, "w", encoding="utf-8") as f:
                f.write(config_content)
        except (PermissionError, OSError) as e:
            log_warning("Could not create raggy_config_example.yaml", e, quiet=False)
            return False
    
    return True


def _print_setup_summary() -> None:
    """Print summary of environment setup completion."""
    print("✅ Environment setup complete!")
    print("\nCreated files:")
    print("- .venv/ (virtual environment)")
    print("- pyproject.toml (project configuration)")
    print("- raggy_config_example.yaml (example configuration)")
    print("- docs/DEVELOPMENT_STATE.md (AI agent continuity tracking)")
    print("\nNext steps:")
    print("1. Add your documentation files to the docs/ directory")
    print("2. Optional: Copy raggy_config_example.yaml to raggy_config.yaml and customize")
    print("3. Run: python raggy.py build")
    print("4. Run: python raggy.py search \"your query\"")


def install_if_missing(packages: List[str], skip_cache: bool = False):
    """Auto-install required packages if missing using uv"""
    if not check_uv_available():
        sys.exit(1)

    # Check if environment is set up properly
    env_ok, env_issue = check_environment_setup()
    if not env_ok:
        if env_issue == "virtual_environment":
            print("ERROR: No virtual environment found.")
            print("Run 'python raggy.py init' to set up the project environment.")
        elif env_issue == "pyproject":
            print("ERROR: No pyproject.toml found.")
            print("Run 'python raggy.py init' to set up the project environment.")
        elif env_issue == "invalid_venv":
            print("ERROR: Invalid virtual environment found.")
            print("Delete .venv directory and run 'python raggy.py init' to recreate it.")
        sys.exit(1)

    # Load cache unless skipped
    cache = {} if skip_cache else load_deps_cache()
    cache_updated = False

    for package_spec in packages:
        # Extract just the package name (before any version specifiers)
        package_name = package_spec.split(">=")[0].split("==")[0].split("[")[0]

        # Check cache first
        if not skip_cache and package_name in cache.get("installed", {}):
            continue

        # Handle special cases for import names
        if package_name == "python-magic-bin":
            import_name = "magic"
        elif package_name == "PyPDF2":
            import_name = "PyPDF2"
        else:
            import_name = package_name.replace("-", "_")

        try:
            # Try to import the module
            spec = importlib.util.find_spec(import_name)
            if spec is None:
                raise ImportError(f"No module named '{import_name}'")

            # Cache successful import
            if "installed" not in cache:
                cache["installed"] = {}
            cache["installed"][package_name] = time.time()
            cache_updated = True

        except ImportError:
            print(f"Installing {package_name}...")
            try:
                # Use uv pip install with the virtual environment
                subprocess.check_call(["uv", "pip", "install", package_spec])
                # Cache successful installation
                if "installed" not in cache:
                    cache["installed"] = {}
                cache["installed"][package_name] = time.time()
                cache_updated = True
            except subprocess.CalledProcessError as e:
                print(f"Failed to install {package_name}: {e}")
                if package_name == "python-magic-bin":
                    print("Trying alternative magic package...")
                    try:
                        subprocess.check_call(["uv", "pip", "install", "python-magic"])
                        cache["installed"][package_name] = time.time()
                        cache_updated = True
                    except subprocess.CalledProcessError:
                        print("Warning: Could not install python-magic. File type detection may be limited.")

    # Save updated cache
    if cache_updated:
        save_deps_cache(cache)


def setup_dependencies(skip_cache: bool = False, quiet: bool = False):
    """Setup dependencies with optional caching"""
    
    # Check if we're in a virtual environment and switch to it if needed
    env_ok, env_issue = check_environment_setup()
    if env_ok:
        # Ensure we're using the virtual environment's Python
        venv_path = Path(".venv")
        if sys.platform == "win32":
            venv_python = venv_path / "Scripts" / "python.exe"
        else:
            venv_python = venv_path / "bin" / "python"
        
        # If we're not already running in the venv, restart with venv python
        if str(venv_python.resolve()) != str(Path(sys.executable).resolve()):
            if not quiet:
                print("Switching to virtual environment...")
            # Re-run the current command with the venv python
            os.execv(str(venv_python), [str(venv_python)] + sys.argv)
    
    required_packages = [
        "chromadb>=0.4.0",
        "sentence-transformers>=2.2.0",
        "PyPDF2>=3.0.0",
        "python-docx>=1.0.0",
    ]

    # Add platform-specific packages
    if sys.platform == "win32":
        required_packages.append("python-magic-bin>=0.4.14")
    else:
        required_packages.append("python-magic")

    if not quiet:
        print("Checking dependencies...")
    install_if_missing(required_packages, skip_cache)

    # Import after installation
    global chromadb, SentenceTransformer, PyPDF2, HAS_MAGIC, magic

    import chromadb
    from sentence_transformers import SentenceTransformer
    import PyPDF2

    # Optional import for file type detection
    try:
        import magic

        HAS_MAGIC = True
    except ImportError:
        HAS_MAGIC = False
        if not quiet:
            print(
                "Note: python-magic not available. Using file extensions for type detection."
            )


class DocumentProcessor:
    """Handles file discovery, text extraction, and chunking operations."""
    
    def __init__(
        self, 
        docs_dir: Path, 
        config: Dict[str, Any],
        quiet: bool = False
    ) -> None:
        self.docs_dir = docs_dir
        self.config = config
        self.quiet = quiet
        
        # File type handlers (Strategy pattern)
        self._file_handlers = {
            ".pdf": self._extract_text_from_pdf,
            ".md": self._extract_text_from_md,
            ".docx": self._extract_text_from_docx,
            ".txt": self._extract_text_from_txt,
        }
    
    def find_documents(self) -> List[Path]:
        """Find all supported documents in docs directory."""
        if not self.docs_dir.exists():
            if not self.quiet:
                print(f"Creating docs directory: {self.docs_dir}")
            self.docs_dir.mkdir(parents=True, exist_ok=True)
            if not self.quiet:
                print(f"Please add your documentation files to {self.docs_dir}")
            return []

        files = []
        for pattern in GLOB_PATTERNS:
            files.extend(self.docs_dir.glob(pattern))
        
        return sorted(files)
    
    def get_file_hash(self, file_path: Path) -> Optional[str]:
        """Public helper to compute file hash for change detection."""
        try:
            return self._get_file_hash(file_path)
        except Exception as error:  # pragma: no cover - defensive guard
            handle_file_error(file_path, "hash", error, quiet=self.quiet)
            return None

    def process_document(
        self, file_path: Path, file_hash: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Process a single document into chunks."""
        if not self.quiet:
            print(f"Processing: {file_path.relative_to(self.docs_dir)}")

        # Validate file path for security
        if not validate_path(file_path, self.docs_dir):
            log_warning(f"Skipping file outside docs directory: {file_path.name}", quiet=self.quiet)
            return []

        # Check file size limits
        try:
            file_size = file_path.stat().st_size
            if file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
                log_warning(f"Skipping large file (>{MAX_FILE_SIZE_MB}MB): {file_path.name}", quiet=self.quiet)
                return []
        except OSError:
            log_warning(f"Could not check file size for {file_path.name}", quiet=self.quiet)
            return []

        try:
            # Extract text using Strategy pattern
            file_extension = file_path.suffix.lower()
            handler = self._file_handlers.get(file_extension)
            
            if handler is None:
                if not self.quiet:
                    supported_types = ', '.join(self._file_handlers.keys())
                    print(f"Skipping unsupported file type: {file_path.name}")
                    print(f"Supported types: {supported_types}")
                return []
            
            text = handler(file_path)

            if not text.strip():
                log_warning(f"No text extracted from {file_path.name}", quiet=self.quiet)
                return []

            # Generate chunks
            chunk_data = self._chunk_text(text)
            chunk_data = self._consolidate_small_chunks(chunk_data)

            # Create document entries
            documents = []
            if file_hash is None:
                file_hash = self._get_file_hash(file_path)

            for i, chunk_info in enumerate(chunk_data):
                doc_id = f"{file_path.stem}_{file_hash[:8]}_{i}"

                # Merge chunk metadata with file metadata
                metadata = {
                    "source": str(file_path.relative_to(self.docs_dir)),
                    "chunk_index": i,
                    "total_chunks": len(chunk_data),
                    "file_hash": file_hash,
                    "file_type": file_path.suffix.lower(),
                    "is_current_state": file_path.name == "CURRENT_STATE.md",
                }
                metadata.update(chunk_info.get("metadata", {}))

                documents.append(
                    {"id": doc_id, "text": chunk_info["text"], "metadata": metadata}
                )

            return documents

        except Exception as e:
            handle_file_error(file_path, "process", e, quiet=self.quiet)
            return []
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Generate SHA256 hash of file for change detection using streaming for large files."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files efficiently
            for chunk in iter(lambda: f.read(CHUNK_READ_SIZE), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _extract_text_template(
        self, file_path: Path, extraction_method: Callable[[Path], str]
    ) -> str:
        """Template method for text extraction with consistent error handling."""
        try:
            result = extraction_method(file_path)
            return result.strip() if result else ""
        except ImportError as e:
            # Handle specific import errors (like missing python-docx)
            library = str(e).split("'")[1] if "'" in str(e) else "dependency"
            warning = f"Warning: {library} not available. Cannot read {file_path.name}"
            print(warning)
            return ""
        except Exception as e:
            sanitized_error = sanitize_error_message(str(e))
            print(f"Warning: Could not extract text from {file_path.name}: {sanitized_error}")
            return ""

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        return self._extract_text_template(file_path, self._extract_pdf_content)

    def _extract_text_from_md(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        return self._extract_text_template(file_path, self._extract_md_content)

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document (.docx)."""
        return self._extract_text_template(file_path, self._extract_docx_content)

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        return self._extract_text_template(file_path, self._extract_txt_content)

    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file."""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            return "\n".join(text_parts)

    def _extract_md_content(self, file_path: Path) -> str:
        """Extract content from Markdown file."""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from Word document."""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)

    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract content from plain text file with encoding fallback."""
        # Try UTF-8 first
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 for older files
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
        smart: bool = True,
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks with optional smart chunking."""
        chunk_size = chunk_size or self.config["search"].get("chunk_size", DEFAULT_CHUNK_SIZE)
        overlap = overlap or self.config["search"].get("chunk_overlap", DEFAULT_CHUNK_OVERLAP)

        if smart and self.config["chunking"]["smart"]:
            return self._chunk_text_smart(text, chunk_size, overlap)
        else:
            return self._chunk_text_simple(text, chunk_size, overlap)

    def _chunk_text_simple(
        self, text: str, chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Simple chunking for backward compatibility."""
        if len(text) <= chunk_size:
            return [{"text": text, "metadata": {"chunk_type": "simple"}}]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + chunk_size - 200, start), -1):
                    if text[i] in ".!?\n":
                        end = i + 1
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    {"text": chunk_text, "metadata": {"chunk_type": "simple"}}
                )

            start = end - overlap

        return chunks

    def _chunk_text_smart(
        self, text: str, base_chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Smart chunking with markdown awareness."""
        chunks = []

        # Split by major sections first (headers)
        sections = re.split(r"(^#{1,6}\s+.*$)", text, flags=re.MULTILINE)
        
        if not self.quiet:
            print(f"  Smart chunking: Found {len(sections)} sections")

        current_header = None
        current_content = ""
        section_count = 0

        for section in sections:
            if re.match(r"^#{1,6}\s+", section):
                # Process previous section if exists
                if current_content.strip():
                    section_count += 1
                    if not self.quiet:
                        print(f"    Processing section {section_count}: {current_header[:50] if current_header else 'No header'}...")
                    section_chunks = self._process_section(
                        current_content, current_header, base_chunk_size, overlap
                    )
                    chunks.extend(section_chunks)

                # Start new section
                current_header = section.strip()
                current_content = ""
            else:
                current_content += section

        # Process final section
        if current_content.strip():
            section_count += 1
            if not self.quiet:
                print(f"    Processing final section {section_count}: {current_header[:50] if current_header else 'No header'}...")
            section_chunks = self._process_section(
                current_content, current_header, base_chunk_size, overlap
            )
            chunks.extend(section_chunks)

        # If no headers found, fall back to simple chunking
        if not chunks:
            if not self.quiet:
                print("  No headers found, using simple chunking")
            return self._chunk_text_simple(text, base_chunk_size, overlap)

        if not self.quiet:
            print(f"  Generated {len(chunks)} chunks")
        return chunks

    def _consolidate_small_chunks(
        self, chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Merge tiny chunks to keep concise documents manageable."""
        if not chunks:
            return chunks

        min_chunk_size = self.config["chunking"].get("min_chunk_size", DEFAULT_CHUNK_SIZE)
        total_chars = sum(len(chunk.get("text", "")) for chunk in chunks)

        # Only consolidate if a short document produced many fragments
        if len(chunks) <= 10 or total_chars == 0:
            return chunks
        if total_chars >= max(min_chunk_size * 3, 1200):
            return chunks

        min_target = max(min_chunk_size // 2, 150)
        merged: List[Dict[str, Any]] = []
        buffer_text = ""
        buffer_meta: Dict[str, Any] = {}

        for chunk in chunks:
            text = chunk.get("text", "").strip()
            if not text:
                continue

            if not buffer_text:
                buffer_text = text
                buffer_meta = dict(chunk.get("metadata", {}))
                continue

            if len(buffer_text) < min_target or len(text) < min_target:
                separator = "\n\n" if buffer_text else ""
                buffer_text = f"{buffer_text}{separator}{text}" if buffer_text else text
                buffer_meta = dict(buffer_meta or chunk.get("metadata", {}))
                buffer_meta["coalesced"] = True
            else:
                merged.append({"text": buffer_text, "metadata": dict(buffer_meta)})
                buffer_text = text
                buffer_meta = dict(chunk.get("metadata", {}))

        if buffer_text:
            merged.append({"text": buffer_text, "metadata": dict(buffer_meta)})

        return merged if merged else chunks

    def _process_section(
        self, content: str, header: Optional[str], chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Process a section with its header."""
        chunks = []
        content = content.strip()

        if not content:
            return chunks

        # Determine chunk size based on content type
        lines = content.split("\n")
        if any(line.strip().startswith(("-", "*", "1.")) for line in lines[:5]):
            # List content - use smaller chunks
            target_size = min(chunk_size, self.config["chunking"]["min_chunk_size"] * 2)
        else:
            # Regular content - use dynamic sizing
            target_size = min(
                max(len(content) // 3, self.config["chunking"]["min_chunk_size"]),
                self.config["chunking"]["max_chunk_size"],
            )

        # Include header in first chunk if preserving headers
        if header and self.config["chunking"]["preserve_headers"]:
            content = f"{header}\n\n{content}"

        # Ensure overlap never eliminates forward progress
        effective_overlap = min(overlap, max(target_size - 1, 0))

        # Split content into chunks
        if len(content) <= target_size:
            metadata = {"chunk_type": "smart"}
            if header:
                metadata["section_header"] = header
                metadata["header_depth"] = len(re.findall(r"^#", header))
            
            chunks.append(
                {
                    "text": content,
                    "metadata": metadata,
                }
            )
        else:
            start = 0
            last_start = -1
            chunk_index = 0

            while start < len(content):
                end = start + target_size

                # Try to break at paragraph or sentence boundary within a bounded window
                if end < len(content):
                    # Look for paragraph breaks first
                    found_boundary = False
                    for i in range(end, max(start + target_size - 300, start), -1):
                        if i - 2 >= start and content[i - 2 : i] == "\n\n":
                            end = i
                            found_boundary = True
                            break
                    if not found_boundary:
                        # Fall back to sentence breaks
                        for i in range(end, max(start + target_size - 200, start), -1):
                            if i < len(content) and content[i] in ".!?\n":
                                end = i + 1
                                break

                # Prevent degenerate range and guarantee progress
                if end <= start:
                    if not self.quiet:
                        print(f"      Warning: Adjusting degenerate chunk window at position {start}; enforcing minimal progress")
                    end = min(start + max(1, target_size // 2), len(content))

                chunk_text = content[start:end].strip()

                if chunk_text:
                    metadata = {
                        "chunk_type": "smart",
                        "section_chunk_index": chunk_index,
                    }
                    if header:
                        metadata["section_header"] = header
                        metadata["header_depth"] = len(re.findall(r"^#", header))
                    
                    chunks.append(
                        {
                            "text": chunk_text,
                            "metadata": metadata,
                        }
                    )
                    chunk_index += 1

                # Compute next start ensuring strict forward progress
                next_start = end - effective_overlap
                if next_start <= start:
                    next_start = start + 1
                last_start = start
                start = next_start

        return chunks


class DatabaseManager:
    """Handles ChromaDB operations and collection management."""
    
    def __init__(
        self,
        db_dir: Path,
        collection_name: str = "project_docs",
        quiet: bool = False
    ) -> None:
        self.db_dir = db_dir
        self.collection_name = collection_name
        self.quiet = quiet
        self._client = None
    
    @property
    def client(self):
        """Lazy-load ChromaDB client."""
        if self._client is None:
            try:
                global chromadb
                self._client = chromadb.PersistentClient(path=str(self.db_dir))
            except NameError:
                # chromadb not imported yet, try importing
                import chromadb
                self._client = chromadb.PersistentClient(path=str(self.db_dir))
        return self._client
    
    def build_index(
        self,
        documents: List[Dict[str, Any]],
        embeddings: Any,
        force_rebuild: bool = False
    ) -> None:
        """Build or update the vector database."""
        try:
            if force_rebuild:
                try:
                    self.client.delete_collection(self.collection_name)
                    if not self.quiet:
                        print("Deleted existing collection")
                except Exception:
                    pass  # Collection may not exist

            collection = self.client.get_or_create_collection(
                name=self.collection_name,
                metadata={"description": "Project documentation embeddings"},
            )
            
            # Clean metadata - remove None values which ChromaDB can't handle
            cleaned_metadatas = []
            for doc in documents:
                cleaned_meta = {k: v for k, v in doc["metadata"].items() if v is not None}
                cleaned_metadatas.append(cleaned_meta)
            
            # Add to ChromaDB
            texts = [doc["text"] for doc in documents]
            collection.add(
                embeddings=embeddings.tolist(),
                documents=texts,
                metadatas=cleaned_metadatas,
                ids=[doc["id"] for doc in documents],
            )
            
        except Exception as e:
            log_error("Failed to build index", e, quiet=self.quiet)
            raise

    def delete_by_ids(self, ids: List[str]) -> None:
        """Delete documents from the collection by id."""
        if not ids:
            return

        try:
            collection = self.get_collection()
            # Chroma can only delete up to a limited batch size reliably; chunk if needed.
            batch_size = 500
            for start in range(0, len(ids), batch_size):
                collection.delete(ids=ids[start : start + batch_size])
        except Exception as error:
            log_warning("Failed to delete documents by id", error, quiet=self.quiet)

    def delete_by_sources(self, sources: List[str]) -> None:
        """Delete all chunks that originated from specific sources."""
        if not sources:
            return

        try:
            collection = self.get_collection()
        except Exception as error:
            log_warning("Could not access collection for deletion", error, quiet=self.quiet)
            return

        for source in sources:
            try:
                collection.delete(where={"source": source})
            except Exception as error:  # pragma: no cover - defensive guard
                log_warning(f"Failed to delete source {source}", error, quiet=self.quiet)

    def get_collection(self):
        """Get the collection for search operations."""
        return self.client.get_collection(self.collection_name)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get database statistics."""
        try:
            collection = self.get_collection()
            count = collection.count()

            # Get source distribution
            all_data = collection.get()
            sources = {}
            for meta in all_data["metadatas"]:
                src = meta["source"]
                sources[src] = sources.get(src, 0) + 1

            return {
                "total_chunks": count,
                "sources": sources,
                "db_path": str(self.db_dir),
            }
        except Exception:
            return {
                "error": "Database not found. Run 'python raggy.py build' first to index your documents."
            }


class SearchEngine:
    """Handles semantic search, hybrid search, and scoring operations."""
    
    def __init__(
        self,
        database_manager: DatabaseManager,
        query_processor: QueryProcessor,
        config: Dict[str, Any],
        manifest_path: Optional[Path] = None,
        quiet: bool = False
    ) -> None:
        self.database_manager = database_manager
        self.query_processor = query_processor
        self.config = config
        self.manifest_path = manifest_path
        self.quiet = quiet
        self._bm25_scorer = None
        self._documents_cache = None
    
    def search(
        self,
        query: str,
        embedding_model: Any,
        n_results: int = DEFAULT_RESULTS,
        hybrid: bool = False,
        expand_query: bool = False,
        show_scores: bool = None,
        path_filter: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """Search the vector database with enhanced capabilities."""
        try:
            collection = self.database_manager.get_collection()
        except Exception:
            log_error("Database collection not found - run 'python raggy.py build' first", quiet=self.quiet)
            return []

        try:
            # Process query
            if expand_query:
                query_info = self.query_processor.process(query)
                processed_query = query_info["processed"]
            else:
                query_info = {
                    "original": query,
                    "type": "keyword",
                    "boost_exact": False,
                }
                processed_query = query

            include_params = ["documents", "metadatas", "distances"]

            manifest = (
                load_manifest(self.manifest_path, quiet=self.quiet)
                if self.manifest_path
                else _default_manifest()
            )
            expected_dim = manifest.get("embedding_dim")
            query_embeddings = None
            use_embeddings = True

            try:
                query_vector = embedding_model.encode([processed_query])
                if hasattr(query_vector, "tolist"):
                    query_embeddings = query_vector.tolist()
                else:  # pragma: no cover - defensive fallback
                    query_embeddings = [list(query_vector)]
            except Exception as error:
                log_warning("Failed to encode query embedding; falling back to text search", error, quiet=self.quiet)
                use_embeddings = False

            if (
                use_embeddings
                and expected_dim
                and query_embeddings
                and len(query_embeddings[0]) != expected_dim
            ):
                log_warning(
                    "Embedding dimension mismatch detected (query vs index); falling back to text search",
                    quiet=self.quiet,
                )
                use_embeddings = False

            # Fetch more results if filtering to ensure we have enough after filtration
            # If path_filter is set, we need a much larger window to find the specific file
            fetch_k = n_results * 50 if path_filter else (n_results * 2 if hybrid else n_results)

            if use_embeddings and query_embeddings:
                results = collection.query(
                    query_embeddings=query_embeddings,
                    n_results=fetch_k,
                    include=include_params,
                )
            else:
                results = collection.query(
                    query_texts=[processed_query],
                    n_results=fetch_k,
                    include=include_params,
                )

            formatted_results = []

            # Initialize BM25 scorer for hybrid search
            if hybrid and self._bm25_scorer is None:
                self._init_bm25_scorer(collection)

            for i in range(len(results["documents"][0])):
                # Apply path filtering if requested
                source = results["metadatas"][0][i].get("source", "")
                if path_filter:
                    if not source.startswith(path_filter):
                        continue

                distance = (
                    results["distances"][0][i] if "distances" in results else None
                )

                # Normalize semantic similarity score
                semantic_score = (
                    normalize_cosine_distance(distance)
                    if distance is not None
                    else 0
                )

                # Calculate keyword score if using hybrid search
                if hybrid and self._bm25_scorer:
                    keyword_score = self._bm25_scorer.score(query, i)
                    # Combine scores
                    final_score = normalize_hybrid_score(
                        semantic_score,
                        keyword_score,
                        self.config["search"]["hybrid_weight"],
                    )
                else:
                    keyword_score = 0
                    final_score = semantic_score

                # Apply exact match boost
                if (
                    query_info.get("boost_exact")
                    and query.lower() in results["documents"][0][i].lower()
                ):
                    final_score = min(1.0, final_score * 1.5)

                # Boost CURRENT_STATE.md content
                if results["metadatas"][0][i].get("is_current_state"):
                    final_score = min(1.0, final_score * 2.0)

                formatted_results.append(
                    {
                        "text": results["documents"][0][i],
                        "metadata": results["metadatas"][0][i],
                        "semantic_score": semantic_score,
                        "keyword_score": keyword_score,
                        "final_score": final_score,
                        "score_interpretation": interpret_score(final_score),
                        "distance": distance,  # Keep for backward compatibility
                        "similarity": final_score,  # Keep for backward compatibility
                    }
                )

            # Sort by final score and limit results
            formatted_results.sort(key=lambda x: x["final_score"], reverse=True)
            formatted_results = formatted_results[:n_results]

            # Rerank results if enabled
            if self.config["search"]["rerank"]:
                formatted_results = self._rerank_results(query, formatted_results)

            # Add highlighting if requested
            show_scores = (
                show_scores
                if show_scores is not None
                else self.config["search"]["show_scores"]
            )
            if show_scores:
                for result in formatted_results:
                    result["highlighted_text"] = self._highlight_matches(
                        query, result["text"]
                    )

            return formatted_results

        except Exception as e:
            log_error("Search error", e, quiet=self.quiet)
            return []

    def _init_bm25_scorer(self, collection):
        """Initialize BM25 scorer with collection documents."""
        if self._documents_cache is None:
            # Get all documents from collection
            all_data = collection.get()
            self._documents_cache = all_data["documents"]

        self._bm25_scorer = BM25Scorer()
        self._bm25_scorer.fit(self._documents_cache)

    def _rerank_results(
        self, query: str, results: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Rerank results to improve diversity and relevance."""
        if len(results) <= 2:
            return results

        reranked = []
        used_sources = set()

        # First pass: take best result from each source
        for result in results:
            source = result["metadata"]["source"]
            if source not in used_sources:
                reranked.append(result)
                used_sources.add(source)
                if len(reranked) >= len(results) // 2:
                    break

        # Second pass: add remaining results
        for result in results:
            if result not in reranked:
                reranked.append(result)

        return reranked[: len(results)]

    def _highlight_matches(
        self, query: str, text: str, context_chars: int = None
    ) -> str:
        """Highlight matching terms in text."""
        context_chars = context_chars or self.config["search"]["context_chars"]

        # Simple highlighting - find first match and show context
        query_terms = re.findall(r"\b\w+\b", query.lower())
        text_lower = text.lower()

        # Find first match position
        match_pos = -1
        for term in query_terms:
            pos = text_lower.find(term)
            if pos != -1:
                match_pos = pos
                break

        if match_pos == -1:
            # No direct match, return beginning
            return text[:context_chars] + "..." if len(text) > context_chars else text

        # Calculate context window around match
        start = max(0, match_pos - context_chars // 2)
        end = min(len(text), match_pos + context_chars // 2)

        # Extend to word boundaries
        while start > 0 and text[start] != " ":
            start -= 1
        while end < len(text) and text[end] != " ":
            end += 1

        excerpt = text[start:end].strip()
        if start > 0:
            excerpt = "..." + excerpt
        if end < len(text):
            excerpt = excerpt + "..."

        return excerpt


class UniversalRAG:
    """Main orchestrator for the RAG system."""
    
    def __init__(
        self,
        docs_dir: str = "./docs",
        db_dir: str = "./vectordb",
        model_name: str = DEFAULT_MODEL,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        quiet: bool = False,
        config_path: Optional[str] = None,
    ) -> None:
        self.docs_dir = Path(docs_dir).resolve()
        self.db_dir = Path(db_dir).resolve()
        self.model_name = model_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.quiet = quiet

        # Load configuration
        self.config = load_config(config_path)

        # Initialize components
        self.document_processor = DocumentProcessor(
            self.docs_dir, self.config, quiet=self.quiet
        )
        self.database_manager = DatabaseManager(
            self.db_dir, quiet=self.quiet
        )
        self.query_processor = QueryProcessor(
            self.config["search"].get("expansions", {})
        )
        self.search_engine = SearchEngine(
            self.database_manager,
            self.query_processor,
            self.config,
            manifest_path=self.db_dir / MANIFEST_FILENAME,
            quiet=self.quiet
        )

        # Lazy-loaded attributes
        self._embedding_model = None
        maintenance_config = self.config.get("maintenance", {})
        thresholds_config = maintenance_config.get("thresholds", {})
        self.thresholds = {
            "soft_chunk_limit": thresholds_config.get(
                "soft_chunk_limit", DEFAULT_SOFT_CHUNK_LIMIT
            ),
            "hard_chunk_limit": thresholds_config.get(
                "hard_chunk_limit", DEFAULT_HARD_CHUNK_LIMIT
            ),
            "soft_document_limit": thresholds_config.get(
                "soft_document_limit", DEFAULT_SOFT_DOCUMENT_LIMIT
            ),
            "per_document_chunk_limit": thresholds_config.get(
                "per_document_chunk_limit", DEFAULT_PER_DOC_CHUNK_LIMIT
            ),
        }

        retention_config = maintenance_config.get("retention", {})
        self.hot_window_days = max(
            0, retention_config.get("hot_days", DEFAULT_HOT_RETENTION_DAYS)
        )
        self.min_hot_updates = max(
            0, retention_config.get("min_hot_updates", DEFAULT_MIN_HOT_UPDATES)
        )

        paths_config = maintenance_config.get("paths", {})
        archive_subdir = paths_config.get("archive_dir", DEFAULT_ARCHIVE_SUBDIR)
        digest_subdir = paths_config.get("digest_dir", DEFAULT_DIGEST_SUBDIR)
        self.archive_dir = (self.docs_dir / Path(archive_subdir)).resolve()
        if not validate_path(self.archive_dir, self.docs_dir):
            self.archive_dir = (self.docs_dir / DEFAULT_ARCHIVE_SUBDIR).resolve()

        self.digest_dir = (self.docs_dir / Path(digest_subdir)).resolve()
        if not validate_path(self.digest_dir, self.docs_dir):
            self.digest_dir = (self.docs_dir / DEFAULT_DIGEST_SUBDIR).resolve()

        auto_config = maintenance_config.get("auto_compact", {})
        self.rebuild_after_compact = bool(
            auto_config.get("rebuild_after_compact", True)
        )

        self._compacting = False

    @property
    def embedding_model(self):
        """Lazy-load embedding model."""
        if self._embedding_model is None:
            if not self.quiet:
                print(f"Loading embedding model ({self.model_name})...")
            self._embedding_model = SentenceTransformer(self.model_name)
        return self._embedding_model

    def build(self, force_rebuild: bool = False, auto_approve: bool = False) -> None:
        """Build or update the vector database with manifest-driven deduplication."""
        start_time = time.time()

        manifest_path = self.db_dir / MANIFEST_FILENAME
        manifest = load_manifest(manifest_path, quiet=self.quiet)
        existing_docs: Dict[str, Any] = manifest.get("documents", {})
        previous_model = manifest.get("model_name")
        model_changed = previous_model and previous_model != self.model_name

        if model_changed and not force_rebuild and not self.quiet:
            print(
                "Embedding model changed from "
                f"{previous_model} → {self.model_name}; performing full rebuild."
            )

        effective_rebuild = force_rebuild or bool(model_changed)

        files = self.document_processor.find_documents()
        if not files:
            log_error("No documents found in docs/ directory", quiet=self.quiet)
            if not self.quiet:
                print("Solution: Add supported files to the docs/ directory")
                print("Supported formats: .md, .pdf, .docx, .txt")
                print("Example: docs/readme.md, docs/guide.pdf, docs/manual.docx, docs/notes.txt")
            return

        if not self.quiet:
            print(f"Found {len(files)} documents")

        file_infos: List[Tuple[Path, str, Optional[str]]] = []
        sources_on_disk: Set[str] = set()
        for file_path in files:
            source = str(file_path.relative_to(self.docs_dir))
            file_hash = self.document_processor.get_file_hash(file_path)
            if file_hash is None:
                continue
            file_infos.append((file_path, source, file_hash))
            sources_on_disk.add(source)

        if not file_infos:
            log_error("No readable documents available for indexing", quiet=self.quiet)
            return

        ids_to_delete: List[str] = []
        removed_sources = sorted(set(existing_docs.keys()) - sources_on_disk)
        for removed_source in removed_sources:
            ids_to_delete.extend(existing_docs.get(removed_source, {}).get("chunk_ids", []))

        all_documents: List[Dict[str, Any]] = []
        manifest_updates: Dict[str, Any] = {}
        unchanged_docs: Dict[str, Any] = {}

        if effective_rebuild and existing_docs and not force_rebuild:
            # If model drift triggered the rebuild we still want a clean slate in the DB
            force_rebuild = True

        for index, (file_path, source, file_hash) in enumerate(file_infos, 1):
            if not self.quiet:
                status = "rebuilding" if effective_rebuild else "checking"
                print(f"[{index}/{len(file_infos)}] {status} {source}")

            manifest_entry = existing_docs.get(source)
            if not effective_rebuild and manifest_entry and manifest_entry.get("file_hash") == file_hash:
                unchanged_docs[source] = manifest_entry
                continue

            if manifest_entry:
                ids_to_delete.extend(manifest_entry.get("chunk_ids", []))

            docs = self.document_processor.process_document(file_path, file_hash=file_hash)
            if not docs:
                continue

            all_documents.extend(docs)
            manifest_updates[source] = {
                "file_hash": file_hash,
                "chunk_ids": [doc["id"] for doc in docs],
                "chunk_count": len(docs),
                "last_indexed": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "source": source,
            }

        if not all_documents and not ids_to_delete and not removed_sources:
            if not self.quiet:
                print("No changes detected; index is already up-to-date.")
            # Ensure manifest still records the current model for future checks
            manifest["model_name"] = self.model_name
            save_manifest(manifest_path, manifest)
            return

        if not effective_rebuild:
            # Remove stale chunks before inserting replacements
            unique_ids = sorted(set(ids_to_delete))
            self.database_manager.delete_by_ids(unique_ids)
            self.database_manager.delete_by_sources(removed_sources)

        if not all_documents and effective_rebuild:
            log_error("No content could be extracted from documents", quiet=self.quiet)
            return

        embeddings = None
        embedding_dim = manifest.get("embedding_dim")
        chunk_total = 0
        if all_documents:
            if not self.quiet:
                print(f"Generating embeddings for {len(all_documents)} chunks...")
            texts = [doc["text"] for doc in all_documents]
            embeddings = self.embedding_model.encode(
                texts, show_progress_bar=not self.quiet
            )
            embedding_dim = len(embeddings[0]) if len(all_documents) else embedding_dim
            chunk_total = len(all_documents)

            self.database_manager.build_index(
                all_documents,
                embeddings,
                force_rebuild=effective_rebuild,
            )

        # Update manifest state combining unchanged + newly processed docs
        updated_manifest_docs = {
            **{src: data for src, data in unchanged_docs.items() if src not in removed_sources},
            **manifest_updates,
        }
        manifest.update(
            {
                "documents": updated_manifest_docs,
                "model_name": self.model_name,
                "embedding_dim": embedding_dim,
                "last_built_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            }
        )
        save_manifest(manifest_path, manifest)

        elapsed = time.time() - start_time
        processed_files = len(manifest_updates)
        unchanged_files = len(unchanged_docs)
        removed_count = len(removed_sources)
        manifest_chunk_total = sum(
            entry.get("chunk_count", 0) for entry in updated_manifest_docs.values()
        )

        print(
            f"{SYMBOLS['success']} Indexed {chunk_total} chunks | "
            f"updated {processed_files} files, unchanged {unchanged_files}, removed {removed_count}"
        )
        print(f"Database saved to: {self.db_dir}")
        if not self.quiet:
            print(f"Build completed in {elapsed:.1f} seconds")

        if not self._compacting:
            self._handle_post_build_maintenance(
                manifest_docs=updated_manifest_docs,
                manifest_chunk_total=manifest_chunk_total,
                auto_approve=auto_approve,
            )
    
    def search(
        self,
        query: str,
        n_results: int = DEFAULT_RESULTS,
        hybrid: bool = False,
        expand_query: bool = False,
        show_scores: bool = None,
        path_filter: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """Search the vector database with enhanced capabilities."""
        if path_filter:
            # Normalize path: remove leading ./ and docs/ prefix if present
            path_filter = path_filter.replace("\\", "/")
            if path_filter.startswith("./"):
                path_filter = path_filter[2:]
            if path_filter.startswith("docs/"):
                path_filter = path_filter[5:]

        return self.search_engine.search(
            query, 
            self.embedding_model,
            n_results, 
            hybrid, 
            expand_query, 
            show_scores,
            path_filter
        )
    
    def interactive_search(self) -> None:
        """Interactive search mode."""
        print(f"\n{SYMBOLS['search']} Interactive Search Mode")
        print("Type your queries (or 'quit' to exit)")
        print("-" * 50)

        while True:
            try:
                query = input("\nQuery: ").strip()
                if query.lower() in ["quit", "exit", "q"]:
                    break

                if not query:
                    continue

                start_time = time.time()
                results = self.search(query)
                elapsed = time.time() - start_time

                if not results:
                    print("No results found.")
                    continue

                print(
                    f"\n{SYMBOLS['found']} Found {len(results)} results (in {elapsed:.3f}s):"
                )
                for i, result in enumerate(results, 1):
                    print(f"\n--- Result {i} ---")
                    print(f"Source: {result['metadata']['source']}")
                    print(
                        f"Chunk: {result['metadata']['chunk_index'] + 1}/{result['metadata']['total_chunks']}"
                    )
                    if result["similarity"]:
                        print(f"Similarity: {result['similarity']:.3f}")
                    print(f"Text preview: {result['text'][:200]}...")

            except KeyboardInterrupt:
                break

        print(f"\n{SYMBOLS['bye']} Goodbye!")
    
    def get_stats(self) -> Dict[str, Any]:
        """Get database statistics."""
        stats = self.database_manager.get_stats()
        if "error" in stats:
            return stats

        manifest_path = self.db_dir / MANIFEST_FILENAME
        manifest = load_manifest(manifest_path, quiet=self.quiet)
        manifest_docs: Dict[str, Any] = manifest.get("documents", {})

        docs_on_disk = []
        try:
            docs_on_disk = [
                str(path.relative_to(self.docs_dir))
                for path in self.document_processor.find_documents()
            ]
        except Exception as error:  # pragma: no cover - defensive guard
            log_warning("Unable to enumerate docs directory for status", error, quiet=self.quiet)

        disk_sources = set(docs_on_disk)
        indexed_sources = set(stats.get("sources", {}).keys())
        manifest_sources = set(manifest_docs.keys())

        missing_from_manifest = sorted(disk_sources - manifest_sources)
        removed_from_disk = sorted(manifest_sources - disk_sources)

        chunk_mismatches = []
        for source in manifest_sources & indexed_sources:
            manifest_count = manifest_docs.get(source, {}).get("chunk_count")
            indexed_count = stats["sources"].get(source)
            if manifest_count is not None and indexed_count is not None and manifest_count != indexed_count:
                chunk_mismatches.append(
                    {
                        "source": source,
                        "manifest": manifest_count,
                        "indexed": indexed_count,
                    }
                )

        stats.update(
            {
                "docs_on_disk": sorted(docs_on_disk),
                "manifest_doc_total": len(manifest_docs),
                "missing_from_manifest": missing_from_manifest,
                "removed_in_manifest": removed_from_disk,
                "chunk_mismatches": chunk_mismatches,
                "embedding_model": manifest.get("model_name") or self.model_name,
                "embedding_dim": manifest.get("embedding_dim"),
                "last_built_at": manifest.get("last_built_at"),
            }
        )

        manifest_chunk_total = sum(
            entry.get("chunk_count", 0) for entry in manifest_docs.values()
        )
        maintenance = self._evaluate_thresholds(
            manifest_docs=manifest_docs,
            manifest_chunk_total=manifest_chunk_total,
            db_chunk_total=stats.get("total_chunks"),
        )

        stats["manifest_chunk_total"] = manifest_chunk_total
        stats["maintenance"] = maintenance
        stats["thresholds"] = maintenance.get("thresholds", {})

        return stats

    def _evaluate_thresholds(
        self,
        *,
        manifest_docs: Dict[str, Any],
        manifest_chunk_total: int,
        db_chunk_total: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Evaluate index size against configured maintenance thresholds."""
        chunk_counts: Dict[str, int] = {
            source: max(0, int(data.get("chunk_count", 0)))
            for source, data in manifest_docs.items()
        }
        doc_total = len(chunk_counts)
        chunk_total_reference = max(
            manifest_chunk_total,
            db_chunk_total if isinstance(db_chunk_total, int) else 0,
        )

        thresholds = self.thresholds
        soft_chunk_limit = thresholds.get("soft_chunk_limit", DEFAULT_SOFT_CHUNK_LIMIT)
        hard_chunk_limit = thresholds.get("hard_chunk_limit", DEFAULT_HARD_CHUNK_LIMIT)
        soft_document_limit = thresholds.get("soft_document_limit", DEFAULT_SOFT_DOCUMENT_LIMIT)
        per_document_limit = thresholds.get(
            "per_document_chunk_limit", DEFAULT_PER_DOC_CHUNK_LIMIT
        )

        per_doc_exceeded: List[Tuple[str, int]] = []
        if per_document_limit:
            per_doc_exceeded = [
                (source, count)
                for source, count in chunk_counts.items()
                if count > per_document_limit
            ]

        largest_source: Optional[str] = None
        largest_count = 0
        if chunk_counts:
            largest_source, largest_count = max(
                chunk_counts.items(), key=lambda item: item[1]
            )

        messages: List[str] = []
        soft_trigger = False
        hard_trigger = False

        if soft_chunk_limit and chunk_total_reference > soft_chunk_limit:
            soft_trigger = True
            messages.append(
                f"Chunk soft limit exceeded: {chunk_total_reference} chunks > {soft_chunk_limit}"
            )

        if soft_document_limit and doc_total > soft_document_limit:
            soft_trigger = True
            messages.append(
                f"Document soft limit exceeded: {doc_total} docs > {soft_document_limit}"
            )

        if hard_chunk_limit and chunk_total_reference >= hard_chunk_limit:
            hard_trigger = True
            messages.append(
                f"Hard chunk limit reached: {chunk_total_reference} chunks ≥ {hard_chunk_limit}"
            )

        if per_doc_exceeded:
            soft_trigger = True
            preview = ", ".join(
                f"{source} ({count})" for source, count in per_doc_exceeded[:5]
            )
            messages.append(
                f"Per-document limit exceeded ({per_document_limit} chunks) by: {preview}"
            )
            if len(per_doc_exceeded) > 5:
                messages.append(
                    f"(+{len(per_doc_exceeded) - 5} additional documents above limit)"
                )

        return {
            "thresholds": thresholds,
            "chunk_total": chunk_total_reference,
            "manifest_chunk_total": manifest_chunk_total,
            "document_total": doc_total,
            "largest_document": {
                "source": largest_source,
                "chunks": largest_count,
            },
            "per_document_exceeded": per_doc_exceeded,
            "soft_limit_exceeded": soft_trigger,
            "hard_limit_exceeded": hard_trigger,
            "messages": messages,
            "needs_attention": soft_trigger or hard_trigger,
        }

    def _handle_post_build_maintenance(
        self,
        *,
        manifest_docs: Dict[str, Any],
        manifest_chunk_total: int,
        auto_approve: bool,
    ) -> None:
        """Check thresholds after build and offer automated compaction."""
        try:
            db_stats = self.database_manager.get_stats()
            db_chunk_total = (
                db_stats.get("total_chunks")
                if isinstance(db_stats, dict)
                else None
            )
        except Exception:
            db_chunk_total = None

        state = self._evaluate_thresholds(
            manifest_docs=manifest_docs,
            manifest_chunk_total=manifest_chunk_total,
            db_chunk_total=db_chunk_total,
        )

        if not state.get("needs_attention"):
            return

        if not self.quiet:
            print("\nMaintenance thresholds:")
            print(
                f"  Chunks: {state['chunk_total']} (manifest {state['manifest_chunk_total']})"
            )
            print(f"  Documents: {state['document_total']}")
            largest = state.get("largest_document", {})
            if largest.get("source"):
                print(
                    f"  Largest source: {largest['source']} ({largest['chunks']} chunks)"
                )
            for message in state.get("messages", []):
                print(f"  - {message}")

        interactive = sys.stdin.isatty() and not self.quiet

        if state.get("hard_limit_exceeded"):
            if auto_approve or interactive:
                if not auto_approve and interactive:
                    answer = input(
                        "Hard limit reached. Run compaction now? [y/N]: "
                    ).strip().lower()
                    if answer not in ("y", "yes"):
                        log_error(
                            "Compaction declined while hard limit is active; build cannot continue.",
                            quiet=self.quiet,
                        )
                        raise RuntimeError("Compaction required but declined")
                summary = self.compact(auto_confirm=True, quiet=self.quiet, reason="hard_limit")
                if not self.quiet:
                    if summary.get("changed"):
                        print(
                            f"{SYMBOLS['success']} Compacted {summary['archived_updates']} updates."
                        )
                    else:
                        print("No compaction changes were required.")
            else:
                log_error(
                    "Hard limit exceeded and no approval available. Rerun with --yes or execute 'python raggy.py compact'.",
                    quiet=self.quiet,
                )
                raise RuntimeError("Compaction required but not approved")
            return

        if state.get("soft_limit_exceeded"):
            if auto_approve:
                summary = self.compact(auto_confirm=True, quiet=self.quiet, reason="soft_limit")
                if not self.quiet and summary.get("changed"):
                    print(
                        f"{SYMBOLS['success']} Compacted {summary['archived_updates']} updates."
                    )
                return

            if interactive:
                answer = input(
                    "Maintenance thresholds exceeded. Run compaction now? [y/N]: "
                ).strip().lower()
                if answer in ("y", "yes"):
                    summary = self.compact(
                        auto_confirm=True,
                        quiet=self.quiet,
                        reason="soft_limit",
                    )
                    if not self.quiet and summary.get("changed"):
                        print(
                            f"{SYMBOLS['success']} Compacted {summary['archived_updates']} updates."
                        )
                else:
                    if not self.quiet:
                        print(
                            "Skipping compaction for now. Run 'python raggy.py compact' when ready."
                        )
            else:
                if not self.quiet:
                    print(
                        "Maintenance thresholds exceeded in non-interactive mode. Run 'python raggy.py compact --yes' to resolve."
                    )

    def compact(
        self,
        *,
        auto_confirm: bool = False,
        quiet: Optional[bool] = None,
        reason: str = "manual",
    ) -> Dict[str, Any]:
        """Run the automated compaction workflow."""
        quiet = self.quiet if quiet is None else quiet
        summary = self._compact_development_state(auto_confirm=auto_confirm, quiet=quiet)

        if summary.get("changed") and self.rebuild_after_compact:
            if not quiet:
                print("Rebuilding index after compaction...")
            self._compacting = True
            try:
                self.build(force_rebuild=False, auto_approve=True)
            finally:
                self._compacting = False
            summary["rebuild_triggered"] = True

        summary["reason"] = reason
        return summary

    def _compact_development_state(
        self, *, auto_confirm: bool, quiet: bool
    ) -> Dict[str, Any]:
        """Archive aged updates from the configured compaction source file into monthly files and digests."""
        summary: Dict[str, Any] = {
            "changed": False,
            "archived_updates": 0,
            "kept_updates": 0,
            "archive_targets": {},
            "digest_updates": {},
        }

        # Get compaction source from config (default: CHANGELOG.md, legacy: DEVELOPMENT_STATE.md)
        compaction_source = self.config.get("maintenance", {}).get("paths", {}).get("compaction_source", "CHANGELOG.md")
        source_path = self.docs_dir / compaction_source
        if not source_path.exists():
            log_error(f"{compaction_source} not found; nothing to compact.", quiet=quiet)
            return summary

        try:
            content = source_path.read_text(encoding="utf-8")
        except Exception as error:
            handle_file_error(source_path, "read", error, quiet=quiet)
            return summary

        update_pattern = re.compile(
            r"^##\s+Update\s+-\s+(\d{4}-\d{2}-\d{2})(?:\s*\(([^\n)]*)\))?\s*\n",
            re.MULTILINE,
        )
        matches = list(update_pattern.finditer(content))
        if not matches:
            if not quiet:
                print(f"No update sections detected in {compaction_source}; skipping compaction.")
            return summary

        updates: List[Dict[str, Any]] = []
        for idx, match in enumerate(matches):
            start = match.start()
            end = matches[idx + 1].start() if idx + 1 < len(matches) else len(content)
            block = content[start:end]
            date_str = match.group(1)
            title = (match.group(2) or "").strip()
            parsed_date = None
            try:
                parsed_date = datetime.strptime(date_str, "%Y-%m-%d").date()
            except ValueError:
                parsed_date = None

            updates.append(
                {
                    "id": idx,
                    "start": start,
                    "end": end,
                    "block": block,
                    "date": parsed_date,
                    "date_str": date_str,
                    "title": title,
                }
            )

        if not updates:
            return summary

        cutoff_date = datetime.now().date() - timedelta(days=self.hot_window_days)

        # Ensure a minimum number of updates stay hot even if older than cutoff
        hot_keep_ids: Set[int] = set()
        dated_updates = [update for update in updates if update["date"] is not None]
        if dated_updates and self.min_hot_updates > 0:
            dated_sorted = sorted(dated_updates, key=lambda item: item["date"], reverse=True)
            for candidate in dated_sorted[: self.min_hot_updates]:
                hot_keep_ids.add(candidate["id"])

        archive_candidates: List[Dict[str, Any]] = []
        keep_candidates: List[Dict[str, Any]] = []

        for update in updates:
            parsed_date = update["date"]
            if parsed_date is None:
                keep_candidates.append(update)
                continue

            if update["id"] in hot_keep_ids:
                keep_candidates.append(update)
                continue

            age_days = (datetime.now().date() - parsed_date).days
            if self.hot_window_days >= 0 and age_days > self.hot_window_days:
                archive_candidates.append(update)
            else:
                keep_candidates.append(update)

        if not archive_candidates:
            if not quiet:
                print("All updates fall within the hot window; no compaction necessary.")
            summary["kept_updates"] = len(updates)
            return summary

        if not auto_confirm:
            if not sys.stdin.isatty() or quiet:
                log_warning(
                    "Compaction requires confirmation; rerun with --yes to proceed.",
                    quiet=quiet,
                )
                summary["kept_updates"] = len(updates)
                return summary

            preview_titles = ", ".join(
                (candidate["title"] or candidate["date_str"]) for candidate in archive_candidates[:3]
            )
            prompt = (
                f"Archive {len(archive_candidates)} updates older than {self.hot_window_days} days"
                f"? ({preview_titles}{'...' if len(archive_candidates) > 3 else ''}) [y/N]: "
            )
            answer = input(prompt).strip().lower()
            if answer not in ("y", "yes"):
                if not quiet:
                    print("Compaction cancelled.")
                summary["kept_updates"] = len(updates)
                return summary

        archive_ids = {candidate["id"] for candidate in archive_candidates}
        slices: List[str] = []
        cursor = 0
        for update in updates:
            if update["id"] in archive_ids:
                slices.append(content[cursor : update["start"]])
                cursor = update["end"]
        slices.append(content[cursor:])
        updated_content = "".join(slices)
        updated_content = updated_content.rstrip() + "\n"

        archive_counts: Dict[str, int] = defaultdict(int)
        digest_counts: Dict[str, int] = defaultdict(int)

        for candidate in archive_candidates:
            archive_date = candidate.get("date")
            if archive_date is None:
                archive_file = self.archive_dir / "undated.md"
                archive_header = "# Development State Archive — Undated\n\n"
                summary_label = "undated"
            else:
                archive_file = self.archive_dir / f"{archive_date.year}-{archive_date.month:02d}.md"
                month_label = archive_date.strftime("%B %Y")
                archive_header = f"# Development State Archive — {month_label}\n\n"
                summary_label = f"{archive_date.year}-{archive_date.month:02d}"

            if not validate_path(archive_file, self.docs_dir):
                log_error(
                    f"Archive path {archive_file} is outside docs/. Skipping this entry.",
                    quiet=quiet,
                )
                continue

            archive_file.parent.mkdir(parents=True, exist_ok=True)
            existing_archive = ""
            if archive_file.exists():
                try:
                    existing_archive = archive_file.read_text(encoding="utf-8")
                except Exception:
                    existing_archive = ""
            else:
                try:
                    archive_file.write_text(archive_header, encoding="utf-8")
                except Exception as error:
                    handle_file_error(archive_file, "write", error, quiet=quiet)
                    continue

            block = candidate["block"].strip()
            if block not in existing_archive:
                try:
                    with archive_file.open("a", encoding="utf-8") as handle:
                        if not existing_archive:
                            # Header already written
                            pass
                        handle.write("\n" + block + "\n\n")
                except Exception as error:
                    handle_file_error(archive_file, "append", error, quiet=quiet)
                    continue

            archive_counts[str(archive_file.relative_to(self.docs_dir))] += 1

            if archive_date is not None:
                digest_file = self.digest_dir / f"{archive_date.year}-{archive_date.month:02d}.md"
                if validate_path(digest_file, self.docs_dir):
                    digest_file.parent.mkdir(parents=True, exist_ok=True)
                    digest_header = (
                        f"# Monthly Digest — {archive_date.strftime('%B %Y')}\n\n"
                    )
                    try:
                        if not digest_file.exists():
                            digest_file.write_text(digest_header, encoding="utf-8")
                    except Exception as error:
                        handle_file_error(digest_file, "write", error, quiet=quiet)
                        continue

                    summary_line = (
                        f"- {candidate['date_str']}: {candidate['title'] or 'Update'} → "
                        f"{archive_file.relative_to(self.docs_dir)}"
                    )
                    try:
                        existing_digest = digest_file.read_text(encoding="utf-8")
                    except Exception:
                        existing_digest = ""
                    if summary_line not in existing_digest:
                        try:
                            with digest_file.open("a", encoding="utf-8") as handle:
                                handle.write(summary_line + "\n")
                        except Exception as error:
                            handle_file_error(digest_file, "append", error, quiet=quiet)
                            continue
                    digest_counts[str(digest_file.relative_to(self.docs_dir))] += 1

        try:
            source_path.write_text(updated_content, encoding="utf-8")
        except Exception as error:
            handle_file_error(source_path, "write", error, quiet=quiet)
            return summary

        summary.update(
            {
                "changed": True,
                "archived_updates": len(archive_candidates),
                "kept_updates": len(keep_candidates),
                "archive_targets": archive_counts,
                "digest_updates": digest_counts,
            }
        )

        if not quiet:
            print(
                f"Archived {len(archive_candidates)} updates into {len(archive_counts)} archive files."
            )

        return summary


    def _get_file_hash(self, file_path: Path) -> str:
        """Generate SHA256 hash of file for change detection using streaming for large files"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files efficiently
            for chunk in iter(lambda: f.read(CHUNK_READ_SIZE), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _extract_text_template(
        self, file_path: Path, extraction_method: Callable[[Path], str]
    ) -> str:
        """Template method for text extraction with consistent error handling."""
        try:
            result = extraction_method(file_path)
            return result.strip() if result else ""
        except ImportError as e:
            # Handle specific import errors (like missing python-docx)
            library = str(e).split("'")[1] if "'" in str(e) else "dependency"
            warning = f"Warning: {library} not available. Cannot read {file_path.name}"
            print(warning)
            return ""
        except Exception as e:
            sanitized_error = sanitize_error_message(str(e))
            print(f"Warning: Could not extract text from {file_path.name}: {sanitized_error}")
            return ""

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        return self._extract_text_template(file_path, self._extract_pdf_content)

    def _extract_text_from_md(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        return self._extract_text_template(file_path, self._extract_md_content)

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document (.docx)."""
        return self._extract_text_template(file_path, self._extract_docx_content)

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        return self._extract_text_template(file_path, self._extract_txt_content)

    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file."""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            return "\n".join(text_parts)

    def _extract_md_content(self, file_path: Path) -> str:
        """Extract content from Markdown file."""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from Word document."""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)

    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract content from plain text file with encoding fallback."""
        # Try UTF-8 first
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 for older files
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()

    def _chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
        smart: bool = True,
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks with optional smart chunking"""
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap

        if smart and self.config["chunking"]["smart"]:
            return self._chunk_text_smart(text, chunk_size, overlap)
        else:
            return self._chunk_text_simple(text, chunk_size, overlap)

    def _chunk_text_simple(
        self, text: str, chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Simple chunking for backward compatibility"""
        if len(text) <= chunk_size:
            return [{"text": text, "metadata": {"chunk_type": "simple"}}]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + chunk_size - 200, start), -1):
                    if text[i] in ".!?\n":
                        end = i + 1
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    {"text": chunk_text, "metadata": {"chunk_type": "simple"}}
                )

            start = end - overlap

        return chunks

    def _chunk_text_smart(
        self, text: str, base_chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Smart chunking with markdown awareness"""
        chunks = []

        # Split by major sections first (headers)
        sections = re.split(r"(^#{1,6}\s+.*$)", text, flags=re.MULTILINE)

        current_header = None
        current_content = ""

        for section in sections:
            if re.match(r"^#{1,6}\s+", section):
                # Process previous section if exists
                if current_content.strip():
                    section_chunks = self._process_section(
                        current_content, current_header, base_chunk_size, overlap
                    )
                    chunks.extend(section_chunks)

                # Start new section
                current_header = section.strip()
                current_content = ""
            else:
                current_content += section

        # Process final section
        if current_content.strip():
            section_chunks = self._process_section(
                current_content, current_header, base_chunk_size, overlap
            )
            chunks.extend(section_chunks)

        # If no headers found, fall back to simple chunking
        if not chunks:
            return self._chunk_text_simple(text, base_chunk_size, overlap)

        return chunks

    def _process_section(
        self, content: str, header: Optional[str], chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Process a section with its header"""
        chunks = []
        content = content.strip()

        if not content:
            return chunks

        # Determine chunk size based on content type
        lines = content.split("\n")
        if any(line.strip().startswith(("-", "*", "1.")) for line in lines[:5]):
            # List content - use smaller chunks
            target_size = min(chunk_size, self.config["chunking"]["min_chunk_size"] * 2)
        else:
            # Regular content - use dynamic sizing
            target_size = min(
                max(len(content) // 3, self.config["chunking"]["min_chunk_size"]),
                self.config["chunking"]["max_chunk_size"],
            )

        # Include header in first chunk if preserving headers
        if header and self.config["chunking"]["preserve_headers"]:
            content = f"{header}\n\n{content}"

        # Ensure overlap never eliminates forward progress
        effective_overlap = min(overlap, max(target_size - 1, 0))

        # Split content into chunks
        if len(content) <= target_size:
            chunks.append(
                {
                    "text": content,
                    "metadata": {
                        "chunk_type": "smart",
                        "section_header": header,
                        "header_depth": len(re.findall(r"^#", header or "")),
                    },
                }
            )
        else:
            start = 0
            last_start = -1
            chunk_index = 0

            while start < len(content):
                end = start + target_size

                # Try to break at paragraph or sentence boundary within a bounded window
                if end < len(content):
                    # Look for paragraph breaks first
                    found_boundary = False
                    for i in range(end, max(start + target_size - 300, start), -1):
                        if i - 2 >= start and content[i - 2 : i] == "\n\n":
                            end = i
                            found_boundary = True
                            break
                    if not found_boundary:
                        # Fall back to sentence breaks
                        for i in range(end, max(start + target_size - 200, start), -1):
                            if i < len(content) and content[i] in ".!?\n":
                                end = i + 1
                                break

                # Prevent degenerate range and guarantee progress
                if end <= start:
                    end = min(start + max(1, target_size // 2), len(content))

                chunk_text = content[start:end].strip()
                if chunk_text:
                    chunks.append(
                        {
                            "text": chunk_text,
                            "metadata": {
                                "chunk_type": "smart",
                                "section_header": header,
                                "header_depth": len(re.findall(r"^#", header or "")),
                                "section_chunk_index": chunk_index,
                            },
                        }
                    )
                    chunk_index += 1

                # Compute next start ensuring strict forward progress
                next_start = end - effective_overlap
                if next_start <= start:
                    next_start = start + 1
                last_start = start
                start = next_start

        return chunks

    def _find_documents(self) -> List[Path]:
        """Find all supported documents in docs directory"""
        if not self.docs_dir.exists():
            print(f"Creating docs directory: {self.docs_dir}")
            self.docs_dir.mkdir(parents=True, exist_ok=True)
            print(f"Please add your documentation files to {self.docs_dir}")
            return []

        patterns = ["**/*.md", "**/*.pdf", "**/*.docx", "**/*.txt"]
        files = []

        for pattern in patterns:
            files.extend(self.docs_dir.glob(pattern))

        return sorted(files)

    def _process_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process a single document into chunks"""
        if not self.quiet:
            print(f"Processing: {file_path.relative_to(self.docs_dir)}")

        # Validate file path for security
        if not validate_path(file_path, self.docs_dir):
            print(f"Warning: Skipping file outside docs directory: {file_path.name}")
            return []

        # Check file size limits (100MB max)
        try:
            file_size = file_path.stat().st_size
            if file_size > 100 * 1024 * 1024:  # 100MB limit
                print(f"Warning: Skipping large file (>100MB): {file_path.name}")
                return []
        except OSError as e:
            print(f"Warning: Could not check file size for {file_path.name}")
            return []

        try:
            # Extract text using Strategy pattern
            file_extension = file_path.suffix.lower()
            handler = self._file_handlers.get(file_extension)
            
            if handler is None:
                if not self.quiet:
                    supported_types = ', '.join(self._file_handlers.keys())
                    print(f"Skipping unsupported file type: {file_path.name}")
                    print(f"Supported types: {supported_types}")
                return []
            
            text = handler(file_path)

            if not text.strip():
                if not self.quiet:
                    print(f"Warning: No text extracted from {file_path}")
                return []

            # Generate chunks
            chunk_data = self._chunk_text(text)

            # Create document entries
            documents = []
            file_hash = self._get_file_hash(file_path)

            for i, chunk_info in enumerate(chunk_data):
                doc_id = f"{file_path.stem}_{file_hash[:8]}_{i}"

                # Merge chunk metadata with file metadata
                metadata = {
                    "source": str(file_path.relative_to(self.docs_dir)),
                    "chunk_index": i,
                    "total_chunks": len(chunk_data),
                    "file_hash": file_hash,
                    "file_type": file_path.suffix.lower(),
                }
                metadata.update(chunk_info.get("metadata", {}))

                documents.append(
                    {"id": doc_id, "text": chunk_info["text"], "metadata": metadata}
                )

            return documents

        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return []


    def run_self_tests(self) -> bool:
        """Run built-in self-tests for raggy functionality"""
        print(f"\n{SYMBOLS['search']} Running raggy self-tests...")
        
        tests_passed = 0
        tests_total = 0
        
        # Test 1: BM25 Scorer
        try:
            print("Testing BM25 scorer...")
            scorer = BM25Scorer()
            test_docs = ["hello world", "world of warcraft", "hello there"]
            scorer.fit(test_docs)
            score = scorer.score("hello world", 0)
            if score > 0:
                print("✓ BM25 scorer working correctly")
                tests_passed += 1
            else:
                print("✗ BM25 scorer test failed")
        except Exception as e:
            print(f"✗ BM25 scorer error: {e}")
        tests_total += 1
        
        # Test 2: Query Processor
        try:
            print("Testing query processor...")
            processor = QueryProcessor()
            result = processor.process("test query")
            if result["original"] == "test query" and "terms" in result:
                print("✓ Query processor working correctly")
                tests_passed += 1
            else:
                print("✗ Query processor test failed")
        except Exception as e:
            print(f"✗ Query processor error: {e}")
        tests_total += 1
        
        # Test 3: Path validation
        try:
            print("Testing path validation...")
            test_path = Path("test.txt")
            is_valid = validate_path(test_path)
            if isinstance(is_valid, bool):
                print("✓ Path validation working correctly")
                tests_passed += 1
            else:
                print("✗ Path validation test failed")
        except Exception as e:
            print(f"✗ Path validation error: {e}")
        tests_total += 1
        
        # Test 4: Scoring normalizer
        try:
            print("Testing scoring normalizer...")
            score = normalize_cosine_distance(0.5)
            interpretation = interpret_score(0.7)
            if 0 <= score <= 1 and interpretation == "Good":
                print("✓ Scoring normalizer working correctly")
                tests_passed += 1
            else:
                print("✗ Scoring normalizer test failed")
        except Exception as e:
            print(f"✗ Scoring normalizer error: {e}")
        tests_total += 1
        
        # Summary
        print(f"\nTest Results: {tests_passed}/{tests_total} tests passed")
        if tests_passed == tests_total:
            print(f"{SYMBOLS['success']} All tests passed!")
            return True
        else:
            print(f"⚠️  {tests_total - tests_passed} tests failed")
            return False

    def diagnose_system(self) -> None:
        """Diagnose system setup and dependencies"""
        print(f"\n{SYMBOLS['search']} Diagnosing raggy system setup...")
        
        # Check Python version
        python_version = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
        print(f"Python version: {python_version}")
        if sys.version_info >= (3, 8):
            print("✓ Python version compatible")
        else:
            print("⚠️  Python 3.8+ recommended")
        
        # Check directories
        print(f"Docs directory: {self.docs_dir}")
        if self.docs_dir.exists():
            doc_count = len(list(self.docs_dir.glob("**/*")))
            print(f"✓ Docs directory exists ({doc_count} files)")
        else:
            print("⚠️  Docs directory not found")
        
        print(f"Database directory: {self.db_dir}")
        if self.db_dir.exists():
            print("✓ Database directory exists")
        else:
            print("ℹ️  Database directory will be created on first build")
        
        # Check dependencies
        print("\nDependency check:")
        deps_status = []
        
        try:
            import chromadb
            print("✓ ChromaDB installed")
            deps_status.append(True)
        except ImportError:
            print("✗ ChromaDB not installed")
            deps_status.append(False)
        
        try:
            from sentence_transformers import SentenceTransformer
            print("✓ sentence-transformers installed")
            deps_status.append(True)
        except ImportError:
            print("✗ sentence-transformers not installed")
            deps_status.append(False)
        
        try:
            import PyPDF2
            print("✓ PyPDF2 installed")
            deps_status.append(True)
        except ImportError:
            print("⚠️  PyPDF2 not installed (PDF support disabled)")
            deps_status.append(False)
        
        try:
            from docx import Document
            print("✓ python-docx installed")
            deps_status.append(True)
        except ImportError:
            print("⚠️  python-docx not installed (DOCX support disabled)")
            deps_status.append(False)
        
        # Model check
        if all(deps_status[:2]):  # ChromaDB and sentence-transformers required
            try:
                print(f"\nTesting embedding model: {self.model_name}")
                model = SentenceTransformer(self.model_name)
                test_embedding = model.encode(["test"])
                print(f"✓ Embedding model loaded successfully (dimensions: {len(test_embedding[0])})")
            except Exception as e:
                print(f"⚠️  Embedding model error: {e}")
        
        # Database status
        try:
            stats = self.get_stats()
            if "error" not in stats:
                print(f"\nDatabase status:")
                print(f"✓ Database accessible")
                print(f"  Total chunks: {stats['total_chunks']}")
                print(f"  Documents indexed: {len(stats['sources'])}")
            else:
                print(f"\nDatabase status:")
                print("ℹ️  No database found - run 'python raggy.py build' to create")
        except Exception as e:
            print(f"⚠️  Database check error: {e}")
        
        print(f"\n{SYMBOLS['success']} Diagnosis complete!")

    def validate_configuration(self) -> bool:
        """Validate configuration and setup"""
        print(f"\n{SYMBOLS['search']} Validating raggy configuration...")
        
        issues = []
        
        # Check config values
        config = self.config
        
        # Validate search config
        search_config = config.get("search", {})
        if not isinstance(search_config.get("hybrid_weight"), (int, float)) or not (0 <= search_config.get("hybrid_weight", 0.7) <= 1):
            issues.append("Invalid hybrid_weight in search config (should be 0.0-1.0)")
        
        if not isinstance(search_config.get("chunk_size"), int) or search_config.get("chunk_size", 1000) < 100:
            issues.append("Invalid chunk_size in search config (should be >= 100)")
        
        if not isinstance(search_config.get("max_results"), int) or search_config.get("max_results", 5) < 1:
            issues.append("Invalid max_results in search config (should be >= 1)")
        
        # Validate chunking config
        chunking_config = config.get("chunking", {})
        min_size = chunking_config.get("min_chunk_size", 300)
        max_size = chunking_config.get("max_chunk_size", 1500)
        
        if not isinstance(min_size, int) or min_size < 50:
            issues.append("Invalid min_chunk_size (should be >= 50)")
        
        if not isinstance(max_size, int) or max_size < min_size:
            issues.append("max_chunk_size should be >= min_chunk_size")
        
        # Check model presets
        models_config = config.get("models", {})
        required_models = ["default", "fast", "multilingual", "accurate"]
        for model_type in required_models:
            if model_type not in models_config:
                issues.append(f"Missing {model_type} model in configuration")
        
        # Validate expansions
        expansions = search_config.get("expansions", {})
        if expansions:
            for term, expansion_list in expansions.items():
                if not isinstance(expansion_list, list) or len(expansion_list) < 2:
                    issues.append(f"Invalid expansion for '{term}' (should be list with original + synonyms)")
        
        # Report results
        if issues:
            print("Configuration issues found:")
            for issue in issues:
                print(f"⚠️  {issue}")
            print(f"\n{len(issues)} issues need attention")
            return False
        else:
            print("✓ Configuration is valid")
            print(f"{SYMBOLS['success']} All validation checks passed!")
            return True


def parse_args() -> Any:
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description="Universal ChromaDB RAG Setup Script v2.0.0 - Enhanced with hybrid search and smart chunking",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Setup:
    %(prog)s init                               # Initialize project environment (first-time setup)
    
  Basic Usage:
    %(prog)s build                              # Build/update index with smart chunking
    %(prog)s search "your search term"         # Semantic search with normalized scores
    %(prog)s status                             # Database statistics and configuration
    %(prog)s compact                            # Archive aged docs and refresh the index
    
  Enhanced Search:
    %(prog)s search "exact phrase" --hybrid    # Hybrid semantic + keyword search
    %(prog)s search "api" --expand             # Query expansion (api → application programming interface)  
    %(prog)s search "documentation" --hybrid --expand # Combined hybrid + expansion
    
  Model Selection:
    %(prog)s build --model-preset multilingual  # Use multilingual model for non-English content
    %(prog)s search "query" --model-preset fast # Quick search with smaller model
    
  Output & Analysis:
    %(prog)s search "query" --json             # Enhanced JSON with score breakdown
    %(prog)s optimize                           # Benchmark semantic vs hybrid search
    %(prog)s interactive --quiet                # Interactive mode, minimal output
    
  Advanced:
    %(prog)s rebuild --config custom.yaml       # Use custom configuration
    %(prog)s search "term" --results 10        # More results with quality scores
        """,
    )

    parser.add_argument(
        "command",
        choices=[
            "init",
            "build",
            "rebuild",
            "search",
            "interactive",
            "status",
            "compact",
            "optimize",
            "test",
            "diagnose",
            "validate",
        ],
        help="Command to execute",
    )
    parser.add_argument("query", nargs="*", help="Search query (for search command)")

    # Options
    parser.add_argument(
        "--docs-dir", default="./docs", help="Documents directory (default: ./docs)"
    )
    parser.add_argument(
        "--db-dir",
        default="./vectordb",
        help="Vector database directory (default: ./vectordb)",
    )
    parser.add_argument(
        "--model", default="all-MiniLM-L6-v2", help="Embedding model name"
    )
    parser.add_argument(
        "--chunk-size", type=int, default=1000, help="Text chunk size (default: 1000)"
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=200,
        help="Text chunk overlap (default: 200)",
    )
    parser.add_argument(
        "--results", type=int, default=5, help="Number of search results (default: 5)"
    )

    # Flags
    parser.add_argument(
        "--fast",
        action="store_true",
        help="Use faster, smaller model (paraphrase-MiniLM-L3-v2)",
    )
    parser.add_argument(
        "--hybrid", action="store_true", help="Use hybrid semantic+keyword search"
    )
    parser.add_argument(
        "--expand", action="store_true", help="Expand query with synonyms"
    )
    parser.add_argument(
        "--path", help="Filter results by file path/directory (e.g., 'docs/CHANGELOG.md')"
    )
    parser.add_argument(
        "--model-preset",
        choices=["fast", "balanced", "multilingual", "accurate"],
        help="Use model preset (overrides --model)",
    )
    parser.add_argument(
        "--skip-deps",
        action="store_true",
        help="Skip dependency checks (faster startup)",
    )
    parser.add_argument("--quiet", "-q", action="store_true", help="Minimal output")
    parser.add_argument(
        "--json", action="store_true", help="Output search results as JSON"
    )
    parser.add_argument(
        "--config", help="Path to config file (default: raggy_config.yaml)"
    )
    parser.add_argument(
        "--yes",
        "-y",
        action="store_true",
        help="Automatically approve maintenance prompts (use with caution)",
    )
    parser.add_argument("--version", action="version", version=f"raggy {__version__}")

    return parser.parse_args()


# Command Pattern Implementation
class Command:
    """Base command interface."""
    
    def execute(self, args: Any, rag: Optional[UniversalRAG] = None) -> None:
        """Execute the command."""
        raise NotImplementedError


class InitCommand(Command):
    """Initialize project environment."""
    
    def execute(self, args: Any, rag: Optional[UniversalRAG] = None) -> None:
        success = setup_environment(quiet=args.quiet)
        if not success:
            sys.exit(1)


class BuildCommand(Command):
    """Build or rebuild the vector database."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        force_rebuild = hasattr(args, 'force_rebuild') and args.force_rebuild
        if hasattr(args, 'command') and args.command == 'rebuild':
            force_rebuild = True
        rag.build(
            force_rebuild=force_rebuild,
            auto_approve=getattr(args, "yes", False),
        )


class SearchCommand(Command):
    """Search the vector database."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        if not args.query:
            log_error("Please provide a search query", quiet=args.quiet)
            return

        query = " ".join(args.query)
        results = rag.search(
            query, 
            n_results=args.results, 
            hybrid=args.hybrid, 
            expand_query=args.expand,
            path_filter=args.path
        )

        if not results:
            print("No results found.")
            return

        if args.json:
            print(
                json.dumps(
                    [
                        {
                            "text": r["text"],
                            "source": r["metadata"]["source"],
                            "chunk": r["metadata"]["chunk_index"] + 1,
                            "final_score": r.get("final_score", r.get("similarity", 0)),
                            "semantic_score": r.get("semantic_score", 0),
                            "keyword_score": r.get("keyword_score", 0),
                            "keyword_score": r.get("keyword_score", 0),
                            "interpretation": r.get("score_interpretation", "Unknown"),
                            "is_current_state": r["metadata"].get("is_current_state", False),
                        }
                        for r in results
                    ],
                    indent=2,
                )
            )
        else:
            print(f"\n{SYMBOLS['search']} Search results for: '{query}'")
            if args.hybrid:
                print("(Using hybrid semantic + keyword search)")
            if args.expand:
                print("(Using query expansion)")
            print("=" * 50)

            for i, result in enumerate(results, 1):
                score = result.get("final_score", result.get("similarity", 0))
                interpretation = result.get("score_interpretation", "")

                score_str = f" ({interpretation}: {score:.3f})" if score else ""
                print(
                    f"\n{i}. {result['metadata']['source']} (chunk {result['metadata']['chunk_index'] + 1}){score_str}"
                )

                # Show highlighted text if available, otherwise truncated text
                display_text = result.get(
                    "highlighted_text", result["text"][:300] + "..."
                )
                print(f"   {display_text}")


class InteractiveCommand(Command):
    """Interactive search mode."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        rag.interactive_search()


class StatusCommand(Command):
    """Show database status and statistics."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        stats = rag.get_stats()
        if "error" in stats:
            # On a fresh setup (including `./raggy init`), the database will not
            # exist yet. Treat this as an informative state, not a hard error.
            print("No database found yet.")
            print("Run './raggy build' to index your documents before requesting status.")
            return

        print(f"Database Statistics:")
        print(f"  Total chunks: {stats['total_chunks']}")
        print(f"  Database path: {stats['db_path']}")
        model_label = stats.get("embedding_model", rag.model_name)
        if model_label != rag.model_name:
            print(f"  Model: {model_label} (configured: {rag.model_name})")
        else:
            print(f"  Model: {model_label}")
        if stats.get("embedding_dim"):
            print(f"  Embedding dim: {stats['embedding_dim']}")
        if stats.get("last_built_at"):
            print(f"  Last build: {stats['last_built_at']}")
        print(f"  Documents on disk: {len(stats.get('docs_on_disk', []))}")
        print(f"  Manifest entries: {stats.get('manifest_doc_total', 0)}")
        print(f"  Indexed documents: {len(stats.get('sources', {}))}")
        print(f"  Config: {'Custom' if args.config else 'Default'}")
        if stats.get("missing_from_manifest"):
            print("  Missing from manifest (needs build):")
            for source in stats["missing_from_manifest"]:
                print(f"    - {source}")
        if stats.get("removed_in_manifest"):
            print("  Orphaned manifest entries (removed on disk):")
            for source in stats["removed_in_manifest"]:
                print(f"    - {source}")
        if stats.get("chunk_mismatches"):
            print("  Chunk mismatches:")
            for mismatch in stats["chunk_mismatches"]:
                print(
                    "    - {source}: manifest={manifest} indexed={indexed}".format(
                        **mismatch
                    )
                )
        print(f"  Documents:")
        for source, count in sorted(stats["sources"].items()):
            print(f"    {source}: {count} chunks")

        maintenance = stats.get("maintenance", {})
        if maintenance:
            thresholds = maintenance.get("thresholds", {})
            print("  Thresholds:")
            print(
                f"    Soft chunks: {thresholds.get('soft_chunk_limit', 'n/a')}"
            )
            print(
                f"    Hard chunks: {thresholds.get('hard_chunk_limit', 'n/a')}"
            )
            print(
                f"    Soft docs: {thresholds.get('soft_document_limit', 'n/a')}"
            )
            print(
                f"    Per-doc limit: {thresholds.get('per_document_chunk_limit', 'n/a')}"
            )

            if maintenance.get("messages"):
                print("  Maintenance warnings:")
                for message in maintenance["messages"]:
                    print(f"    - {message}")
            else:
                print("  Maintenance warnings: none")


class CompactCommand(Command):
    """Archive aged documentation and rebuild the index."""

    def execute(self, args: Any, rag: UniversalRAG) -> None:
        summary = rag.compact(
            auto_confirm=getattr(args, "yes", False),
            quiet=args.quiet,
            reason="manual",
        )

        if summary.get("changed"):
            archived = summary.get("archived_updates", 0)
            targets = ", ".join(summary.get("archive_targets", {}).keys())
            if not args.quiet:
                print(
                    f"{SYMBOLS['success']} Archived {archived} updates. Files: {targets or 'n/a'}"
                )
        else:
            if not args.quiet:
                print("No compaction changes required.")


class OptimizeCommand(Command):
    """Benchmark and optimize search settings."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        print(
            f"\n{SYMBOLS['search']} Running benchmark queries to optimize settings..."
        )

        # Get sample content from the database to generate test queries
        stats = rag.get_stats()
        if "error" in stats or stats["total_chunks"] == 0:
            print("Error: No indexed content found. Run 'build' first.")
            return

        # Generate universal test queries based on document names and common terms
        test_queries = []
        doc_names = list(stats["sources"].keys())

        # Add filename-based queries (without extensions)
        for doc in doc_names[:3]:  # Use first 3 documents
            base_name = Path(doc).stem.replace("-", " ").replace("_", " ")
            test_queries.append(base_name)

        # Add some universal technical queries
        universal_queries = [
            "configuration",
            "setup",
            "guide",
            "documentation",
            "features",
        ]
        test_queries.extend(universal_queries[:2])  # Add 2 universal terms

        print(
            f"Testing with queries derived from your content: {', '.join(test_queries)}"
        )

        results_semantic = []
        results_hybrid = []

        for query in test_queries:
            print(f"Testing: {query}")

            # Test semantic search
            sem_results = rag.search(query, n_results=3, hybrid=False)
            avg_sem_score = (
                sum(r.get("final_score", 0) for r in sem_results) / len(sem_results)
                if sem_results
                else 0
            )
            results_semantic.append(avg_sem_score)

            # Test hybrid search
            hyb_results = rag.search(query, n_results=3, hybrid=True)
            avg_hyb_score = (
                sum(r.get("final_score", 0) for r in hyb_results) / len(hyb_results)
                if hyb_results
                else 0
            )
            results_hybrid.append(avg_hyb_score)

        avg_semantic = sum(results_semantic) / len(results_semantic)
        avg_hybrid = sum(results_hybrid) / len(results_hybrid)

        print(f"\n{SYMBOLS['found']} Optimization Results:")
        print(f"  Average Semantic Score: {avg_semantic:.3f}")
        print(f"  Average Hybrid Score: {avg_hybrid:.3f}")

        if avg_hybrid > avg_semantic * 1.1:
            print(
                f"\n{SYMBOLS['success']} Recommendation: Use --hybrid flag for better results"
            )
        elif avg_semantic > avg_hybrid * 1.1:
            print(
                f"\n{SYMBOLS['success']} Recommendation: Semantic search performs best for this content"
            )
        else:
            print(f"\n{SYMBOLS['success']} Both search modes perform similarly")

        print(f"\nSuggested usage:")
        print(
            f'  python {sys.argv[0]} search \\"your query\\" --hybrid    # For exact matches'
        )
        print(
            f'  python {sys.argv[0]} search \\"your query\\" --expand    # For broader results'
        )


class TestCommand(Command):
    """Run built-in self-tests."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        success = rag.run_self_tests()
        if not success:
            sys.exit(1)


class DiagnoseCommand(Command):
    """Diagnose system setup and dependencies."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        rag.diagnose_system()


class ValidateCommand(Command):
    """Validate configuration and setup."""
    
    def execute(self, args: Any, rag: UniversalRAG) -> None:
        success = rag.validate_configuration()
        if not success:
            sys.exit(1)


class CommandFactory:
    """Factory for creating command instances."""
    
    _commands = {
        "init": InitCommand,
        "build": BuildCommand,
        "rebuild": BuildCommand,
        "search": SearchCommand,
        "interactive": InteractiveCommand,
        "status": StatusCommand,
        "compact": CompactCommand,
        "optimize": OptimizeCommand,
        "test": TestCommand,
        "diagnose": DiagnoseCommand,
        "validate": ValidateCommand,
    }
    
    @classmethod
    def create_command(cls, command_name: str) -> Command:
        """Create a command instance."""
        command_class = cls._commands.get(command_name)
        if command_class is None:
            raise ValueError(f"Unknown command: {command_name}")
        return command_class()


def main() -> None:
    """Main entry point using Command pattern."""
    args = parse_args()

    # Check for updates early (non-intrusive, once per session)
    try:
        config = load_config(args.config) if hasattr(args, 'config') else {}
        check_for_updates(quiet=args.quiet, config=config)
    except Exception:
        pass  # Silently fail - don't interrupt user workflow

    # Create and execute command
    try:
        command = CommandFactory.create_command(args.command)
        
        # Handle init command specially (no RAG instance needed)
        if args.command == "init":
            command.execute(args)
            return

        # Setup dependencies for other commands
        if not args.skip_deps:
            setup_dependencies(quiet=args.quiet)
        else:
            # Still need to import even if skipping dependency checks
            try:
                global chromadb, SentenceTransformer, PyPDF2, HAS_MAGIC, magic
                import chromadb
                from sentence_transformers import SentenceTransformer
                import PyPDF2

                try:
                    import magic
                    HAS_MAGIC = True
                except ImportError:
                    HAS_MAGIC = False
            except ImportError as e:
                log_error(f"Missing dependency: {e}", quiet=args.quiet)
                log_error("Run without --skip-deps or install dependencies manually", quiet=args.quiet)
                return

        # Determine model to use
        model_name = _determine_model(args)

        # Initialize RAG system
        rag = UniversalRAG(
            docs_dir=args.docs_dir,
            db_dir=args.db_dir,
            model_name=model_name,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            quiet=args.quiet,
            config_path=args.config,
        )

        # Execute the command
        command.execute(args, rag)
        
    except ValueError as e:
        log_error(str(e), quiet=args.quiet)
        sys.exit(1)
    except Exception as e:
        log_error(f"Unexpected error executing command '{args.command}'", e, quiet=args.quiet)
        sys.exit(1)


def _determine_model(args: Any) -> str:
    """Determine which model to use based on arguments."""
    if args.model_preset:
        config = load_config(args.config)
        preset_models = {
            "fast": config["models"]["fast"],
            "multilingual": config["models"]["multilingual"],
            "accurate": config["models"]["accurate"],
        }
        return preset_models.get(args.model_preset, config["models"]["default"])
    else:
        return FAST_MODEL if args.fast else args.model


if __name__ == "__main__":
    main()
